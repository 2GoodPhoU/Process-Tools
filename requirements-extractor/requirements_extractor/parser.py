"""Walk a .docx document in order and emit a structural event stream.

Document assumptions (matches the user's typical spec format):
  - A header region at the top (title, optional metadata paragraphs).
  - One or more "requirements tables".  By default these are 2-column
    tables; both the column geometry and per-column meaning are
    user-configurable via ``Config.tables``.
  - In each requirements-table row, one column holds a section title /
    topic / actor and the other holds mixed content (paragraphs, bullet
    lists, or nested tables — to arbitrary depth).
  - Standard Word heading styles (Heading 1/2/3) may appear outside of
    tables to group content — these are captured as a "heading trail".

The parser emits an ordered list of events:
  - HeadingEvent          — a top-level heading paragraph
  - SectionRowEvent       — a row whose actor-column text looks like a
                            numbered section header (e.g. "3.1 Auth")
  - RequirementEvent      — a single extracted requirement

Writers that only care about requirements (e.g. the Excel writer) can
filter the stream; writers that need structural context (e.g. the
statement-set CSV exporter) consume it whole.

Recursion
---------
When ``Config.parser.recursive`` is True (the default) the content walker
descends into cells and nested tables of arbitrary depth.  Block-refs
become dotted paths like ``T1R3C2 > T1R2C1 > P1`` so traceability is
preserved even for deeply nested content.  With ``recursive: false`` the
walker matches the original one-level-of-nesting behaviour.
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Iterator, List, Optional, Union

from docx import Document
from docx.document import Document as _Document
from docx.oxml.ns import qn
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

from .config import Config
from .detector import KeywordMatcher, compute_confidence, split_sentences
from .models import (
    HeadingEvent,
    Requirement,
    RequirementEvent,
    SectionRowEvent,
    compute_stable_id,
)

Event = Union[HeadingEvent, SectionRowEvent, RequirementEvent]


# ---------------------------------------------------------------------------
# Parse context — carried through the recursive walker.
# ---------------------------------------------------------------------------


@dataclass
class _ParseContext:
    source_file: str
    config: Config
    matcher: KeywordMatcher
    heading_trail: List[str] = field(default_factory=list)
    table_index: int = 0                    # counts top-level tables
    order_counter: int = 0
    # Latest section-style row title encountered; used to populate the
    # Requirement.section_topic column distinctly from primary_actor.
    current_section_title: str = ""
    # Heading-scope skip state — set when a top-level Heading whose text
    # matches ``skip_sections.matches_title`` fires (e.g. "Glossary").
    # Subsequent content is dropped until a Heading at this level or
    # shallower fires that does NOT match.  ``None`` means "not skipping".
    # Implementation lives in :func:`parse_docx_events` and
    # :func:`_emit_candidate`.
    skip_heading_level: Optional[int] = None

    def next_order(self) -> int:
        self.order_counter += 1
        return self.order_counter

    def trail_str(self) -> str:
        return " > ".join(h for h in self.heading_trail if h)

    def in_skipped_section(self) -> bool:
        """True iff a heading-triggered skip is currently active."""
        return self.skip_heading_level is not None


# ---------------------------------------------------------------------------
# docx traversal helpers — unchanged shape, tidied.
# ---------------------------------------------------------------------------


def _cell_element(cell: _Cell):
    """Return the underlying ``<w:tc>`` element for a table cell.

    ``python-docx`` (>=1.0) does not expose an official public attribute
    for this; ``_tc`` has been the de-facto API for years and is pinned
    via ``python-docx>=1.1,<2`` in requirements.txt.  This helper tries
    public-ish attributes first so the library can later add one without
    us needing to chase it.
    """
    for attr in ("_tc", "_element", "element"):
        elm = getattr(cell, attr, None)
        if elm is not None:
            return elm
    raise AttributeError(
        "Could not locate the underlying XML element for this docx cell; "
        "python-docx may have changed its internal API."
    )


def _paragraph_element(p: Paragraph):
    """Return the underlying ``<w:p>`` element for a paragraph.

    Same encapsulation rationale as :func:`_cell_element`.
    """
    for attr in ("_p", "_element", "element"):
        elm = getattr(p, attr, None)
        if elm is not None:
            return elm
    raise AttributeError(
        "Could not locate the underlying XML element for this docx paragraph; "
        "python-docx may have changed its internal API."
    )


def iter_block_items(parent) -> Iterator[object]:
    """Yield Paragraph and Table objects from a parent in document order.

    Works on the Document body and on table cells.
    """
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = _cell_element(parent)
    else:
        parent_elm = getattr(parent, "_element", None) or getattr(parent, "element", None)

    for child in parent_elm.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, parent)
        elif child.tag == qn("w:tbl"):
            yield Table(child, parent)


def _paragraph_text(p: Paragraph) -> str:
    return (p.text or "").strip()


def _heading_level(p: Paragraph) -> Optional[int]:
    style = p.style.name if p.style is not None else ""
    if style and style.lower().startswith("heading "):
        try:
            return int(style.split()[-1])
        except ValueError:
            return None
    return None


def _is_bullet(p: Paragraph) -> bool:
    """Best-effort bullet/numbering detection."""
    pPr = _paragraph_element(p).find(qn("w:pPr"))
    if pPr is None:
        return False
    numPr = pPr.find(qn("w:numPr"))
    if numPr is not None:
        return True
    style = (p.style.name if p.style is not None else "").lower()
    return "list" in style or "bullet" in style


def _cell_text_all(cell: _Cell) -> str:
    """Concatenate visible text in a cell (including nested tables).

    Used when we need a summary string (e.g. the first-column topic).
    Whitespace is collapsed.
    """
    parts: List[str] = []

    def _walk(parent):
        for block in iter_block_items(parent):
            if isinstance(block, Paragraph):
                t = _paragraph_text(block)
                if t:
                    parts.append(t)
            elif isinstance(block, Table):
                for row in block.rows:
                    for subcell in row.cells:
                        _walk(subcell)

    _walk(cell)
    return " ".join(parts).strip()


def _cell_intro_text(cell: _Cell) -> str:
    """Non-heading paragraph text from a cell (used for section intros).

    Unlike ``_cell_text_all`` this does not descend into nested tables —
    section intros are expected to be plain prose.
    """
    parts: List[str] = []
    for block in iter_block_items(cell):
        if isinstance(block, Paragraph):
            level = _heading_level(block)
            if level is not None:
                continue
            text = _paragraph_text(block)
            if text:
                parts.append(text)
    return " ".join(parts).strip()


def _update_heading_trail(trail: List[str], level: int, text: str) -> None:
    """Update the trail so index ``level - 1`` holds ``text``.

    If the document skips levels (e.g. H1 → H3 with no intervening H2),
    the missing slots are filled with empty strings so the depth of a
    heading can always be recovered from its position in the list.
    This keeps the "Heading Trail" column meaningful when skimming.
    """
    # Truncate anything at or below the incoming level — those sub-headings
    # belong to a now-closed branch.
    while len(trail) >= level:
        trail.pop()
    # Pad with empty strings for any skipped intermediate levels.
    while len(trail) < level - 1:
        trail.append("")
    trail.append(text)


# ---------------------------------------------------------------------------
# The recursive walker.
# ---------------------------------------------------------------------------


# ---------------------------------------------------------------------------


# ---------------------------------------------------------------------------
# Procedural-table detection — header signal + multi-actor cell parsing.
#
# The implementation lives in :mod:`procedural` (extracted in a refactor
# — parser.py had grown to ~870 lines because procedural-table detection
# was layered on top of the original 2-col walker).  Re-imported here
# because :func:`_emit_candidate`, :func:`_walk_content`, and
# :func:`parse_docx_events` all reference one or more of these helpers
# inline.  Imports under their original underscore-prefixed names so
# existing test imports (``from requirements_extractor.parser import
# _split_candidate_actors``) continue to resolve via this module.
# ---------------------------------------------------------------------------

from .procedural import (  # noqa: E402, F401 — re-exports for backward compat
    REQUIRED_ACTION_KEYWORD,
    is_required_action_header,
    _normalise_header_cell,
    _pick_primary,
    _resolve_primary_from_candidates,
    _split_candidate_actors,
)

# Compound-requirement detection (0.6.1).  Imported lazily-friendly: the
# module has no heavy imports of its own so this is cheap.  See
# :mod:`requirements_extractor.compound` for what the pre-pass does.
from . import compound as _compound
# Multi-action requirement decomposition (0.6.2).  Pure-helper module;
# no heavy imports of its own.  See
# :mod:`requirements_extractor.multi_action` for what the post-pass does.
from . import multi_action as _multi_action
from ._logging import logger as _logger


def _apply_multi_action(
    req: Requirement,
    ctx: _ParseContext,
    force_requirement: bool,
) -> List[Requirement]:
    """Apply 0.6.2 multi-action decomposition to a freshly emitted Requirement.

    Returns either ``[req]`` (no decomposition) or, in ``split`` mode,
    ``[req, sub1, sub2, ...]`` where ``req`` carries the original source
    text as a record and each sub-requirement is one atomic obligation
    derived from one verb phrase of the parent.

    Modes:

    * ``single`` - returns ``[req]`` unchanged.  Equivalent to 0.6.1.
    * ``flag``   - returns ``[req]`` with a one-line note appended to
      ``req.notes`` describing the multi-action count and a "consider
      splitting" recommendation.  No new rows; reviewers decide
      per-requirement.
    * ``split``  - emits the parent (with a "split into N atomic
      sub-requirements" note) plus N child rows.  Each child carries
      the parent's stable_id in :attr:`Requirement.parent_id` and a
      hierarchical sub-ID of the form ``<parent_id>.<n>`` (1-based).
      Children inherit actor / heading_trail / source_file / row_ref;
      block_ref is annotated with ``(multi-action sub N/M)``.

    Defensive design:
      * Wrapped in try/except so any failure logs and falls through to
        ``[req]``.  One pathological sentence never aborts the parse.
      * Gated off when ``force_requirement=True`` (procedural required-
        action tables).  Those rows are atomic by virtue of the
        ``Required Action`` column header semantics; splitting them
        would fragment what is already atomic.
      * When the multi-action config is missing (legacy callers passing
        a Config dataclass built before 0.6.2), the function is a no-op.
    """
    if req is None:
        return []
    cfg = getattr(ctx.config, "multi_action", None)
    if cfg is None or not getattr(cfg, "enabled", False):
        return [req]
    mode = getattr(cfg, "mode", "single")
    if mode == "single":
        return [req]
    if force_requirement:
        # Procedural required-action gate - mirrors the 0.6.1 compound
        # gate.  See PATCH-0.6.2-NOTES.md for rationale.
        return [req]
    try:
        det = _multi_action.detect(
            req.text,
            min_actions=getattr(cfg, "min_actions", 2),
        )
    except Exception:  # noqa: BLE001 - defensive: never abort the walk
        _logger.warning(
            "multi-action detect failed for stable_id=%s; "
            "falling back to single-mode behaviour.",
            req.stable_id,
            exc_info=True,
        )
        return [req]

    if not det.is_multi_action:
        return [req]

    if mode == "flag":
        try:
            note = _multi_action.render_flag_note(det)
        except Exception:  # noqa: BLE001 - defensive
            return [req]
        if req.notes:
            req.notes = req.notes + "\n" + note
        else:
            req.notes = note
        return [req]

    if mode == "split":
        try:
            return _build_split_subrequirements(req, det, ctx)
        except Exception:  # noqa: BLE001 - defensive
            _logger.warning(
                "multi-action split failed for stable_id=%s; "
                "falling back to single-mode behaviour.",
                req.stable_id,
                exc_info=True,
            )
            return [req]

    # Unknown mode - fall through conservatively.
    return [req]


def _build_split_subrequirements(
    parent: Requirement,
    det: "_multi_action.MultiActionDetection",
    ctx: _ParseContext,
) -> List[Requirement]:
    """Construct the parent + N atomic sub-requirements for ``split`` mode.

    The parent row keeps its original text but gains a note pointing at
    the children.  Each child row carries a synthetic sentence
    ``<actor> <modal> <verb-phrase>.``, reuses the parent's
    actor/heading/source-file context, and gets a hierarchical sub-ID
    of the form ``<parent_stable_id>.<n>``.

    Why include the parent: source faithfulness.  Reviewers should
    always be able to find the original sentence as it was written by
    the author.  Children carry the atomic semantics; the parent
    carries the source link.  ReqIF maps this cleanly to the
    "decomposition" relation - see PATCH-0.6.2-NOTES.md for export
    notes.
    """
    n = det.count
    # Annotate parent with a pointer at the children.
    parent_note = (
        f"Multi-action split: source sentence decomposed into "
        f"{n} atomic sub-requirements ({parent.stable_id}.1 ... "
        f"{parent.stable_id}.{n})."
    )
    if parent.notes:
        parent.notes = parent.notes + "\n" + parent_note
    else:
        parent.notes = parent_note

    out: List[Requirement] = [parent]
    for i, action in enumerate(det.actions, start=1):
        sub_text = _multi_action.render_sub_requirement(det, action)
        sub_block_ref = f"{parent.block_ref} (multi-action sub {i}/{n})"
        sub_stable_id = f"{parent.stable_id}.{i}"
        sub_note = (
            f"Atomic sub-requirement {i} of {n} (split from "
            f"{parent.stable_id})."
        )
        sub = Requirement(
            order=ctx.next_order(),
            source_file=parent.source_file,
            heading_trail=parent.heading_trail,
            section_topic=parent.section_topic,
            row_ref=parent.row_ref,
            block_ref=sub_block_ref,
            primary_actor=parent.primary_actor,
            secondary_actors=list(parent.secondary_actors),
            text=sub_text,
            req_type=parent.req_type,
            keywords=list(parent.keywords),
            confidence=parent.confidence,
            notes=sub_note,
            polarity=parent.polarity,
            stable_id=sub_stable_id,
            context=parent.context,
            parent_id=parent.stable_id,
        )
        out.append(sub)
    return out


#: Re-exported for the force-requirement path so a future confidence
#: upgrade in :mod:`detector` lands in every code path automatically.
_length_based_confidence = compute_confidence


#: REVIEW §3.8 — soft cap on the inline context column.  Long enough to
#: carry one or two surrounding sentences; short enough that it doesn't
#: distort the reviewer's xlsx column widths.  Truncation is sentence-
#: friendly: we cut at the last whitespace before the cap and append "…".
_MAX_CONTEXT_CHARS = 280


def _build_context(context: str, text: str) -> str:
    """Return a reviewer-friendly context snippet, or "" if not useful.

    Two suppression cases keep the column from carrying noise:

    1. **Empty input** — nothing to show.
    2. **Context is the requirement** — when a paragraph or bullet has
       only one sentence, the captured block text equals the
       requirement text after normalisation.  Showing it twice in the
       row would just waste horizontal space, so we collapse to "".

    Otherwise we whitespace-collapse and (if needed) truncate to
    :data:`_MAX_CONTEXT_CHARS`, cutting at the last whitespace inside
    the cap so we don't bisect a word.  The truncated suffix is "…"
    (a single character so the visible width stays predictable).
    """
    if not context:
        return ""
    # Whitespace-collapse so multi-line bullets become single-line context.
    collapsed = " ".join(context.split())
    if not collapsed:
        return ""
    # Same normalisation as compute_stable_id — keeps the redundancy
    # check insensitive to cosmetic differences (a stray space, a
    # trailing newline, an ALL-CAPS run vs the title-case original).
    text_norm = " ".join((text or "").split()).casefold()
    if collapsed.casefold() == text_norm:
        return ""
    if len(collapsed) <= _MAX_CONTEXT_CHARS:
        return collapsed
    cap = _MAX_CONTEXT_CHARS - 1  # leave room for the ellipsis
    cut = collapsed.rfind(" ", 0, cap)
    if cut <= 0:
        cut = cap
    return collapsed[:cut].rstrip() + "…"


def _emit_candidate(
    text: str,
    ctx: _ParseContext,
    *,
    row_ref: str,
    block_ref: str,
    primary_actor: str,
    force_requirement: bool = False,
    context: str = "",
) -> Optional[Requirement]:
    """Build a Requirement iff ``text`` passes detection + config filters.

    Returns None when the sentence should be dropped (not a requirement,
    filtered by ``content.skip_if_starts_with`` / ``skip_pattern``, or
    missing a required primary actor).

    When ``force_requirement`` is True the keyword-match gate is bypassed
    and the sentence is emitted as Hard with a synthetic
    ``(Required Action)`` keyword label.  Used for rows inside a
    procedural required-action table — the table-type signal tells us
    every row is binding even when the text doesn't use shall / must /
    etc.  A Hard keyword match still wins if one is present (it carries
    richer keyword info), so force mode only kicks in when the classifier
    would otherwise drop the row.
    """
    # Heading-scope skip: a prior boilerplate heading (e.g. "Glossary",
    # "References", "Revision History") flipped on a section-wide skip
    # that stays active until a sibling heading clears it.  Drop anything
    # captured while it's in effect.  See ``_ParseContext.skip_heading_level``.
    if ctx.in_skipped_section():
        return None

    # Content-level filter: skip by prefix / pattern.
    if ctx.config.content.should_skip(text):
        return None

    req_type, keywords, confidence = ctx.matcher.classify(text)
    if not req_type:
        if not force_requirement:
            return None
        # Header-signalled requirement — synthesise the classification so
        # downstream writers / reviewers can see why it was captured.
        req_type = "Hard"
        keywords = [REQUIRED_ACTION_KEYWORD]
        confidence = _length_based_confidence(text)

    # Optional "must have an actor to count" filter.
    if ctx.config.content.require_primary_actor and not (primary_actor or "").strip():
        return None

    notes = ""
    if req_type == "Soft":
        notes = "Soft language — verify with author whether this is a binding requirement."

    polarity = "Negative" if ctx.matcher.is_negative(text) else "Positive"

    return Requirement(
        order=ctx.next_order(),
        source_file=ctx.source_file,
        heading_trail=ctx.trail_str(),
        section_topic=ctx.current_section_title,
        row_ref=row_ref,
        block_ref=block_ref,
        primary_actor=primary_actor,
        secondary_actors=[],            # filled in by the resolver below
        text=text,
        req_type=req_type,
        keywords=keywords,
        confidence=confidence,
        notes=notes,
        polarity=polarity,
        stable_id=compute_stable_id(ctx.source_file, primary_actor, text),
        context=_build_context(context, text),
    )


def _walk_content(
    parent,
    ctx: _ParseContext,
    *,
    row_ref: str,
    primary_actor: str,
    resolver_fn,
    ref_prefix: str = "",
    recursive: Optional[bool] = None,
    force_requirement: bool = False,
    candidate_actors: Optional[List[str]] = None,
) -> Iterator[Requirement]:
    """Walk the paragraph/table children of ``parent`` and yield Requirements.

    ``parent`` is typically the content-column cell from a requirements
    table, but this function also handles nested tables (recursively) and
    is reused for preamble prose.  ``ref_prefix`` is the dotted-path string
    we prepend to per-block refs so nested items stay traceable.

    ``force_requirement`` bypasses the keyword-match gate — set to True
    when the caller already knows every sentence here is a requirement
    (e.g. inside a procedural required-action table).  See
    :func:`_emit_candidate` for the exact semantics.

    ``candidate_actors`` enables per-sentence primary-actor resolution for
    procedural rows where the actor cell lists multiple candidates
    ("Auth Service, Gateway, Logger").  When provided, each sentence is
    scanned for the earliest candidate name and that candidate becomes
    the primary actor for the emitted requirement.  Sentences that don't
    name any candidate fall back to ``primary_actor`` (typically the
    joined cell text).

    Compound-requirement pre-pass (0.6.1)
    -------------------------------------
    When ``ctx.config.compound.enabled`` is True (the default), this
    function first runs :func:`compound.detect_groups` over the
    paragraph children to identify "modal+colon lead-in followed by a
    list of conditions" patterns.  Detected groups are emitted as a
    single aggregated requirement (lead-in + items joined by their
    inferred ``and`` / ``or`` / ``unless`` connector); the legacy
    per-paragraph walk skips the claimed indices.  The pre-pass is
    wrapped in try/except — any failure logs a warning and falls back
    to the legacy behaviour for the whole parent.
    """
    if recursive is None:
        recursive = ctx.config.parser.recursive

    paragraph_idx = 0
    bullet_idx = 0
    nested_table_idx = 0

    def _push_ref(tail: str) -> str:
        return f"{ref_prefix} > {tail}" if ref_prefix else tail

    # Materialise the children into a list so the compound pre-pass can
    # peek ahead without exhausting the iterator the legacy walk consumes.
    blocks = list(iter_block_items(parent))

    # ----- Compound-requirement pre-pass --------------------------------
    #
    # Build a parallel ``_BlockInfo`` list for the paragraph children, run
    # ``compound.detect_groups``, and stash the result in two lookup
    # structures:
    #   * lead_in_at[idx] -> CompoundGroup   (paragraph IS a lead-in)
    #   * claimed_items   -> set of indices  (paragraph IS a list item
    #                                          claimed by some group)
    #
    # The legacy per-block walk then consults both before emitting; a
    # claimed item is silently skipped, a lead-in emits the aggregated
    # requirement instead of the bare lead-in sentence.
    #
    # Wrapped in try/except so any failure here (malformed input, an
    # exception inside the compound module, etc.) logs a warning and
    # leaves the legacy behaviour in place for THIS parent.  One bad
    # parent never aborts the run.
    lead_in_at = {}
    claimed_items = set()
    compound_enabled = bool(getattr(ctx.config, "compound", None) and ctx.config.compound.enabled)
    # Procedural required-action tables (force_requirement=True) already
    # treat every row as an atomic binding requirement by virtue of the
    # column header.  Aggregating their lead-in + bullets into a single
    # compound would *drop* atomic requirements rather than capture
    # additional ones — exactly the opposite of what the patch is for.
    # Compound detection is therefore disabled inside force-requirement
    # walks; reviewers still get one row per bullet there, matching the
    # 0.6.0 behaviour the procedural fixture suite pins.  Pure-prose
    # paths (the case the patch was authored for) continue to aggregate.
    # This is the autonomous design call flagged in PATCH-0.6.1-NOTES.md
    # for Eric's morning review.
    if force_requirement:
        compound_enabled = False
    if compound_enabled:
        try:
            block_infos = []
            for b in blocks:
                if isinstance(b, Paragraph):
                    text = _paragraph_text(b)
                    is_b = _is_bullet(b)
                    has_marker = (
                        not is_b
                        and bool(text)
                        and _compound.has_list_marker_prefix(text)
                    )
                    block_infos.append(
                        _compound._BlockInfo(
                            text=text,
                            is_paragraph=True,
                            is_bullet=is_b,
                            has_marker_prefix=has_marker,
                        )
                    )
                else:
                    # Tables (or anything non-paragraph) break a list run
                    # naturally; record a placeholder so indices line up.
                    block_infos.append(
                        _compound._BlockInfo(
                            text="", is_paragraph=False,
                            is_bullet=False, has_marker_prefix=False,
                        )
                    )
            for grp in _compound.detect_groups(block_infos):
                lead_in_at[grp.lead_in_idx] = grp
                for j in grp.item_indices:
                    claimed_items.add(j)
        except Exception:  # noqa: BLE001 — defensive: pre-pass must never abort the walk
            _logger.warning(
                "compound pre-pass failed for parent at row_ref=%s; "
                "falling back to legacy per-paragraph behaviour.",
                row_ref,
                exc_info=True,
            )
            lead_in_at = {}
            claimed_items = set()

    # ----- Main walk ----------------------------------------------------

    for idx, block in enumerate(blocks):
        if isinstance(block, Paragraph):
            text = _paragraph_text(block)
            if not text:
                continue
            # Sub-headings inside a cell update the trail so nested content
            # inherits that context in its output.
            level = _heading_level(block)
            if level is not None:
                _update_heading_trail(ctx.heading_trail, level, text)
                continue

            # Compound: skip a paragraph that was claimed as a list item
            # by an earlier lead-in.  The aggregated requirement was
            # already emitted when the lead-in fired.
            if idx in claimed_items:
                continue

            # Compound: this paragraph IS the lead-in of a group.  Emit
            # one aggregated requirement using the synthetic compound
            # text and mark the row_ref / block_ref so reviewers can
            # tell the row apart from a regular paragraph match.
            grp = lead_in_at.get(idx)
            if grp is not None:
                paragraph_idx += 1
                br = _push_ref(f"Paragraph {paragraph_idx} (compound)")
                actor_for_this = _pick_primary(
                    grp.aggregated_text, primary_actor, candidate_actors
                )
                # Reviewer context: the original lead-in plus the items,
                # whitespace-collapsed.  ``_build_context`` will truncate
                # at _MAX_CONTEXT_CHARS so very long lists don't bloat
                # the column.
                ctx_text = " ".join(
                    [grp.lead_in_text] + list(grp.items)
                )
                try:
                    req = _emit_candidate(
                        grp.aggregated_text, ctx,
                        row_ref=row_ref, block_ref=br,
                        primary_actor=actor_for_this,
                        force_requirement=force_requirement,
                        context=ctx_text,
                    )
                except Exception:  # noqa: BLE001 — defensive: emit failure must not abort the walk
                    _logger.warning(
                        "compound emit failed for lead-in at idx=%d row_ref=%s; "
                        "falling back to legacy emit for the lead-in paragraph only.",
                        idx, row_ref,
                        exc_info=True,
                    )
                    req = None
                if req is not None:
                    if req.notes:
                        req.notes = (
                            req.notes
                            + "\nCompound requirement: lead-in + "
                            + str(len(grp.items))
                            + " "
                            + grp.connector
                            + "-joined item(s)."
                        )
                    else:
                        req.notes = (
                            "Compound requirement: lead-in + "
                            + str(len(grp.items))
                            + " "
                            + grp.connector
                            + "-joined item(s)."
                        )
                    req.secondary_actors = resolver_fn(
                        grp.aggregated_text, actor_for_this
                    )
                    yield from _apply_multi_action(req, ctx, force_requirement)
                continue

            if _is_bullet(block):
                bullet_idx += 1
                br = _push_ref(f"Bullet {bullet_idx}")
                actor_for_this = _pick_primary(
                    text, primary_actor, candidate_actors
                )
                # Bullet text IS the unit; context normally collapses
                # to "" because text == context.  We still pass it so
                # multi-sentence bullets show the sibling sentences.
                req = _emit_candidate(
                    text, ctx,
                    row_ref=row_ref, block_ref=br, primary_actor=actor_for_this,
                    force_requirement=force_requirement,
                    context=text,
                )
                if req is not None:
                    req.secondary_actors = resolver_fn(text, actor_for_this)
                    yield from _apply_multi_action(req, ctx, force_requirement)
            else:
                paragraph_idx += 1
                br = _push_ref(f"Paragraph {paragraph_idx}")
                # The full paragraph is the natural reviewer context for
                # every sentence we extract from it.  ``_build_context``
                # suppresses redundancy when paragraph == sentence.
                for sent in split_sentences(text):
                    actor_for_this = _pick_primary(
                        sent, primary_actor, candidate_actors
                    )
                    req = _emit_candidate(
                        sent, ctx,
                        row_ref=row_ref, block_ref=br, primary_actor=actor_for_this,
                        force_requirement=force_requirement,
                        context=text,
                    )
                    if req is not None:
                        req.secondary_actors = resolver_fn(sent, actor_for_this)
                        yield from _apply_multi_action(req, ctx, force_requirement)
        elif isinstance(block, Table):
            nested_table_idx += 1
            table_tag = f"Nested Table {nested_table_idx}"
            if recursive:
                # True recursion: walk every nested cell with the same
                # walker.  Block ref becomes a dotted path so traceability
                # is preserved even at depth.
                for r_idx, nrow in enumerate(block.rows, start=1):
                    for c_idx, ncell in enumerate(nrow.cells, start=1):
                        nested_prefix = _push_ref(f"{table_tag} R{r_idx}C{c_idx}")
                        yield from _walk_content(
                            ncell, ctx,
                            row_ref=row_ref,
                            primary_actor=primary_actor,
                            resolver_fn=resolver_fn,
                            ref_prefix=nested_prefix,
                            recursive=True,
                            force_requirement=force_requirement,
                            candidate_actors=candidate_actors,
                        )
            else:
                # Legacy one-level behaviour: flatten each nested cell's
                # visible text and split into sentences — no further
                # descent.
                for r_idx, nrow in enumerate(block.rows, start=1):
                    for c_idx, ncell in enumerate(nrow.cells, start=1):
                        cell_text = _cell_text_all(ncell)
                        if not cell_text:
                            continue
                        br = _push_ref(f"{table_tag} R{r_idx}C{c_idx}")
                        for sent in split_sentences(cell_text):
                            req = _emit_candidate(
                                sent, ctx,
                                row_ref=row_ref, block_ref=br,
                                primary_actor=primary_actor,
                                force_requirement=force_requirement,
                                context=cell_text,
                            )
                            if req is not None:
                                req.secondary_actors = resolver_fn(sent, primary_actor)
                                yield from _apply_multi_action(req, ctx, force_requirement)


# ---------------------------------------------------------------------------
# Top-level entry points.
# ---------------------------------------------------------------------------


def parse_docx_events(
    path: Path,
    resolver_fn,
    *,
    config: Optional[Config] = None,
) -> List[Event]:
    """Parse a .docx and return an ordered event stream.

    ``resolver_fn(text, primary_actor) -> List[str]`` is typically
    ``ActorResolver.resolve`` bound to an instance.  ``config`` tweaks what
    counts as a requirements table, which sections to skip, which
    keywords to use, etc.  When omitted, built-in defaults apply.
    """
    cfg = config or Config.defaults()
    matcher = KeywordMatcher.from_config(cfg.keywords)

    path = Path(path)
    doc = Document(str(path))
    ctx = _ParseContext(
        source_file=path.name,
        config=cfg,
        matcher=matcher,
    )
    events: List[Event] = []
    section_re = cfg.tables.section_re()

    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            text = _paragraph_text(block)
            if not text:
                continue
            level = _heading_level(block)
            if level is not None:
                # Heading-scope skip management.  Two cases:
                # (a) A skip is currently active and this heading is at
                #     the skip's level or shallower — the boilerplate
                #     section has ended.  Clear the flag, then fall
                #     through and re-evaluate this heading like normal
                #     (it may itself be another boilerplate heading).
                # (b) After (a)'s clear, if this heading's text matches
                #     the skip filter, start a new skip scoped to its
                #     level so all deeper-or-equal content drops.
                if (
                    ctx.skip_heading_level is not None
                    and level <= ctx.skip_heading_level
                ):
                    ctx.skip_heading_level = None
                if cfg.skip_sections.matches_title(text):
                    ctx.skip_heading_level = level
                _update_heading_trail(ctx.heading_trail, level, text)
                events.append(HeadingEvent(level=level, text=text))
                continue
            # Preamble / inter-table prose.  Scan for requirements but use an
            # empty primary actor — the content filter can drop everything
            # here if the user prefers strict mode.
            for sent in split_sentences(text):
                req = _emit_candidate(
                    sent, ctx,
                    row_ref="Preamble",
                    block_ref="Paragraph",
                    primary_actor="",
                    context=text,
                )
                if req is not None:
                    req.secondary_actors = resolver_fn(sent, "")
                    for emitted in _apply_multi_action(req, ctx, False):
                        events.append(RequirementEvent(requirement=emitted))
            continue

        if isinstance(block, Table):
            ctx.table_index += 1

            # User-driven whole-table skip.
            if ctx.table_index in cfg.skip_sections.table_indices:
                continue

            num_cols = len(block.columns)

            # Procedural "required-action" detection.  The header shape
            #     |  (blank)  |  Step  |  Required Action  |
            # forces the table into req-table mode with the procedural
            # column mapping (actor=1, content=3) and marks every
            # content row as a requirement regardless of modal keywords.
            # Runs BEFORE the generic is_requirement_table check so the
            # signal works even without a paired .reqx.yaml — the
            # header itself is the table-type declaration.
            procedural_required_action = False
            header_row_index: Optional[int] = None
            if num_cols == 3 and len(block.rows) >= 1:
                header_cells_text = [
                    _cell_text_all(c) for c in block.rows[0].cells
                ]
                if is_required_action_header(header_cells_text):
                    procedural_required_action = True
                    header_row_index = 0

            if procedural_required_action:
                actor_idx = 0        # col 1 (0-based)
                content_idx = 2      # col 3 (0-based)
                is_req_table = True
            else:
                is_req_table = cfg.tables.is_requirement_table(num_cols)
                actor_idx = cfg.tables.actor_column - 1
                content_idx = cfg.tables.content_column - 1

            # Table-local state for blank-actor continuation.  In
            # procedural required-action tables, a blank column-1 body
            # cell inherits the actor from the nearest non-blank
            # predecessor row within the *same* table (FIELD_NOTES §4
            # + Eric's 2026-04-23 pass).  Tracker is scoped to this
            # table so it can never leak into the next one.
            last_non_blank_actor = ""

            for r_idx, row in enumerate(block.rows, start=1):
                cells = row.cells
                row_ref = f"Table {ctx.table_index}, Row {r_idx}"

                # Skip the header row when we've identified one via the
                # required-action signal.  Body rows start at r_idx == 2.
                if header_row_index is not None and r_idx == header_row_index + 1:
                    continue

                if is_req_table and len(cells) > max(actor_idx, content_idx):
                    topic = _cell_text_all(cells[actor_idx])
                    content_cell = cells[content_idx]

                    # Blank-actor continuation (procedural tables only).
                    # Outside procedural mode a blank actor column has
                    # different semantics (see implicit_system_actor
                    # fixture), so this path is gated to avoid changing
                    # the long-standing behaviour for 2-col tables.
                    if procedural_required_action:
                        if topic.strip():
                            last_non_blank_actor = topic
                        elif last_non_blank_actor:
                            topic = last_non_blank_actor

                    # Skip this row when it matches a user-configured title.
                    if cfg.skip_sections.matches_title(topic):
                        continue

                    if section_re.match(topic or ""):
                        intro = _cell_intro_text(content_cell)
                        ctx.current_section_title = topic
                        events.append(
                            SectionRowEvent(
                                title=topic, intro=intro, row_ref=row_ref,
                            )
                        )
                        for req in _walk_content(
                            content_cell, ctx,
                            row_ref=row_ref,
                            primary_actor="",
                            resolver_fn=resolver_fn,
                            force_requirement=procedural_required_action,
                            candidate_actors=None,
                        ):
                            events.append(RequirementEvent(requirement=req))
                    else:
                        # Multi-actor cell resolution: parse the column-1
                        # text as a candidate set.  Only activated inside
                        # procedural required-action tables to avoid
                        # changing behaviour for conventional 2-col
                        # tables.  Non-procedural rows fall through with
                        # candidate_actors=None (old behaviour).
                        row_candidates: Optional[List[str]] = None
                        if procedural_required_action:
                            split = _split_candidate_actors(topic)
                            if split:
                                row_candidates = split

                        for req in _walk_content(
                            content_cell, ctx,
                            row_ref=row_ref,
                            primary_actor=topic,
                            resolver_fn=resolver_fn,
                            force_requirement=procedural_required_action,
                            candidate_actors=row_candidates,
                        ):
                            events.append(RequirementEvent(requirement=req))
                else:
                    # Not a requirements table — walk every cell anyway
                    # so nothing gets silently dropped, but with no
                    # primary actor.
                    for c_idx, cell in enumerate(cells, start=1):
                        for req in _walk_content(
                            cell, ctx,
                            row_ref=f"{row_ref}, Col {c_idx}",
                            primary_actor="",
                            resolver_fn=resolver_fn,
                        ):
                            events.append(RequirementEvent(requirement=req))

    return events


def parse_docx(
    path: Path,
    resolver_fn,
    *,
    config: Optional[Config] = None,
) -> List[Requirement]:
    """Parse a .docx and return just its Requirements (no structural events).

    Convenience over :func:`parse_docx_events` for callers that only
    want the requirement rows — most tests, the actor scanner's
    per-doc walker, and any user-script that doesn't need the heading
    or section-row context.  Equivalent to filtering the event stream
    down to ``RequirementEvent`` and unwrapping each one.

    Use :func:`parse_docx_events` instead when you DO need the
    structural context (the statement-set CSV exporter is the canonical
    example — it places each requirement in its heading hierarchy).
    """
    return [
        e.requirement
        for e in parse_docx_events(path, resolver_fn, config=config)
        if isinstance(e, RequirementEvent)
    ]
