"""Actor-only extraction mode.

Walks one or more .docx files and harvests *actor observations* —
primary-column cell text from every requirements-table row, plus any
regex / NER hits in the content column.  The observations are grouped
into canonical names (with aliases) and written to an Excel workbook
whose first sheet has the exact ``Actor`` / ``Aliases`` header shape
that :func:`requirements_extractor.actors.load_actors_from_xlsx`
expects.

The output therefore round-trips: run it once on a corpus, tidy the
result, then feed the same file back in via ``--actors`` on a normal
requirements run.

The design goal is complementary rather than duplicative: this mode
deliberately does *not* run the keyword detector — it's cheaper, and
more importantly it also discovers actors from rows that happen not to
contain any "shall" sentences.

Grouping rules (no seed):
  1. Normalise each observation (strip leading determiner, trailing
     possessive, collapse whitespace, lowercase for comparison).
  2. Bucket by the normalised key.
  3. Within a bucket, pick the canonical spelling as the most frequent
     raw form, preferring Title-case, then shorter strings, then
     alphabetic order.
  4. Every other spelling becomes an alias.

Grouping rules (with seed actors.xlsx):
  1. Each seed actor reserves a canonical group up front (its user-typed
     name wins regardless of observation frequency).
  2. Any observation whose normalised form matches the seed's canonical
     or any seed alias folds into that group.
  3. Observed spellings that aren't already seed aliases become new
     aliases on the group.
  4. Seeded actors with zero observations still appear in the output so
     the user's curated list is preserved.
"""

from __future__ import annotations

from collections import Counter
from dataclasses import dataclass, field
from pathlib import Path
from typing import Callable, Dict, List, Optional, Sequence, Tuple

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

from .actors import ActorEntry, ActorResolver, load_actors_from_xlsx
from .config import Config, resolve_config
from .parser import (
    _cell_text_all,
    _heading_level,
    _paragraph_text,
    _update_heading_trail,
    iter_block_items,
)


# ---------------------------------------------------------------------------
# Cancellation — the scan-side counterpart to ExtractionCancelled.
# ---------------------------------------------------------------------------


class ActorScanCancelled(RuntimeError):
    """Raised when ``cancel_check`` returns True mid-scan.

    Mirrors ``ExtractionCancelled`` — a cancelled scan writes no output
    file so there is no half-finished artefact on disk.
    """


# ---------------------------------------------------------------------------
# Data model.
# ---------------------------------------------------------------------------


@dataclass
class ActorObservation:
    """A single raw sighting of an actor-like string inside a document."""

    raw: str                     # exactly as captured, original casing preserved
    normalised: str              # lowercased, determiner-stripped key used for grouping
    source: str                  # "primary" | "regex" | "nlp"
    file: str                    # just the filename
    row_ref: str                 # e.g. "Table 2, Row 4"
    heading_trail: str           # best-effort breadcrumbs


@dataclass
class ActorGroup:
    """A canonical actor + every variant we observed for it."""

    canonical: str
    aliases: List[str]
    sources: List[str]           # sorted union of contributing sources
    count: int                   # total observations across all variants
    files: List[str]             # sorted unique filenames
    first_seen: Optional[str]    # "filename — Row Ref" of first observation
    seeded: bool = False         # True when this group came from the seed list


@dataclass
class ActorScanStats:
    files_processed: int = 0
    observations: int = 0
    groups: int = 0
    errors: List[str] = field(default_factory=list)


@dataclass
class ActorScanResult:
    groups: List[ActorGroup]
    observations: List[ActorObservation]
    stats: ActorScanStats
    output_path: Optional[Path] = None


# ---------------------------------------------------------------------------
# Normalisation — exposed so tests can pin down the contract.
# ---------------------------------------------------------------------------


_DETERMINERS = ("the ", "a ", "an ")
_POSSESSIVE_TAILS = ("\u2019s", "'s", "\u2019", "'")


def normalise_actor_text(text: str) -> str:
    """Return a canonical lookup key for ``text``.

    Strips leading determiners, trailing possessives, collapses whitespace,
    and lowercases.  Returns ``""`` for empty/whitespace input.
    """
    if not text:
        return ""
    s = " ".join(text.split())
    if not s:
        return ""
    low = s.lower()
    for det in _DETERMINERS:
        if low.startswith(det):
            s = s[len(det):]
            low = s.lower()
            break
    for tail in _POSSESSIVE_TAILS:
        if s.endswith(tail):
            s = s[: -len(tail)]
            break
    return s.strip().lower()


# ---------------------------------------------------------------------------
# Walker — collects observations from a single .docx.
# ---------------------------------------------------------------------------


def _walk_one_doc(
    path: Path,
    cfg: Config,
    resolver: Optional[ActorResolver],
) -> List[ActorObservation]:
    """Return every actor observation in ``path``.

    Honours the same config knobs as the main extractor (skip_sections,
    tables geometry, section_prefix) so the two modes agree on which
    cells are "actor cells".
    """
    observations: List[ActorObservation] = []
    doc = Document(str(path))
    heading_trail: List[str] = []
    section_re = cfg.tables.section_re()
    table_index = 0

    actor_idx = cfg.tables.actor_column - 1
    content_idx = cfg.tables.content_column - 1

    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            level = _heading_level(block)
            if level is not None:
                _update_heading_trail(heading_trail, level, _paragraph_text(block))
            continue

        if not isinstance(block, Table):
            continue

        table_index += 1
        if table_index in cfg.skip_sections.table_indices:
            continue
        if not cfg.tables.is_requirement_table(len(block.columns)):
            continue

        trail_str = " > ".join(h for h in heading_trail if h)

        for r_idx, row in enumerate(block.rows, start=1):
            cells = row.cells
            if len(cells) <= max(actor_idx, content_idx):
                continue

            topic = _cell_text_all(cells[actor_idx])
            row_ref = f"Table {table_index}, Row {r_idx}"

            if cfg.skip_sections.matches_title(topic):
                continue
            if not topic:
                continue
            if section_re.match(topic):
                # Section-header row — a structural marker, not an actor.
                continue

            observations.append(
                ActorObservation(
                    raw=topic,
                    normalised=normalise_actor_text(topic),
                    source="primary",
                    file=path.name,
                    row_ref=row_ref,
                    heading_trail=trail_str,
                )
            )

            if resolver is not None:
                content_text = _cell_text_all(cells[content_idx])
                for hit, src in _resolver_hits(resolver, content_text, topic):
                    observations.append(
                        ActorObservation(
                            raw=hit,
                            normalised=normalise_actor_text(hit),
                            source=src,
                            file=path.name,
                            row_ref=row_ref,
                            heading_trail=trail_str,
                        )
                    )

    return observations


def _resolver_hits(
    resolver: ActorResolver, text: str, primary: str,
) -> List[Tuple[str, str]]:
    """Return (name, source) pairs, where source is 'regex' or 'nlp'.

    ``ActorResolver.resolve`` mushes the two together; for scanning we
    want the attribution, so we re-walk the same passes here.  This uses
    the resolver's pre-built alias regex to stay consistent.
    """
    hits: List[Tuple[str, str]] = []
    if not text:
        return hits
    primary_lower = (primary or "").strip().lower()
    seen: set[str] = set()

    if resolver._actor_re is not None:  # noqa: SLF001 — internal re-use is fine
        for m in resolver._actor_re.finditer(text):
            canonical = resolver._alias_to_canonical[m.group(1).lower()]
            key = canonical.lower()
            if key == primary_lower or key in seen:
                continue
            seen.add(key)
            hits.append((canonical, "regex"))

    if resolver._nlp is not None:  # noqa: SLF001
        try:
            ner_doc = resolver._nlp(text)
            for ent in ner_doc.ents:
                if ent.label_ not in {"PERSON", "ORG", "NORP", "PRODUCT"}:
                    continue
                name = ent.text.strip()
                key = name.lower()
                if not name or key == primary_lower or key in seen:
                    continue
                seen.add(key)
                hits.append((name, "nlp"))
        except Exception:  # noqa: BLE001 — NER is best-effort
            pass

    return hits


# ---------------------------------------------------------------------------
# Grouping — pure; operates on observations + optional seeds.
# ---------------------------------------------------------------------------


def _pick_canonical(raws: Counter[str]) -> str:
    """Pick the best canonical spelling from a frequency table.

    Ranking: (frequency, has_upper, -length, lexicographic).
    """
    def key(item: Tuple[str, int]) -> Tuple[int, int, int, str]:
        raw, count = item
        has_upper = 1 if any(c.isupper() for c in raw) else 0
        return (count, has_upper, -len(raw), raw)

    return max(raws.items(), key=key)[0]


def group_observations(
    observations: Sequence[ActorObservation],
    *,
    seed_entries: Optional[Sequence[ActorEntry]] = None,
) -> List[ActorGroup]:
    """Bucket observations into ActorGroups, optionally seeded by known actors."""
    seed_canonical_of: Dict[str, str] = {}
    seed_aliases_of: Dict[str, List[str]] = {}
    seed_names_in_order: List[str] = []
    for entry in seed_entries or []:
        name = (entry.name or "").strip()
        if not name:
            continue
        key = normalise_actor_text(name)
        if not key:
            continue
        seed_canonical_of[key] = name
        seed_aliases_of[name] = [a for a in entry.aliases if a]
        seed_names_in_order.append(name)
        for alias in entry.aliases:
            ak = normalise_actor_text(alias)
            if ak:
                seed_canonical_of[ak] = name

    # bucket_key → accumulator dict
    buckets: Dict[str, Dict] = {}

    def _ensure(key: str, *, seeded_name: Optional[str]) -> Dict:
        b = buckets.get(key)
        if b is None:
            b = {
                "raws": Counter(),
                "sources": set(),
                "files": [],
                "count": 0,
                "first_file": None,
                "first_row_ref": None,
                "seeded_name": seeded_name,
            }
            buckets[key] = b
        return b

    for obs in observations:
        norm = obs.normalised
        if not norm:
            continue
        seeded = seed_canonical_of.get(norm)
        key = seeded if seeded is not None else norm
        b = _ensure(key, seeded_name=seeded)
        b["raws"][obs.raw] += 1
        b["sources"].add(obs.source)
        if obs.file and obs.file not in b["files"]:
            b["files"].append(obs.file)
        b["count"] += 1
        if b["first_file"] is None:
            b["first_file"] = obs.file
            b["first_row_ref"] = obs.row_ref

    # Seeded actors with zero observations should still show up so the
    # user's curated list survives a rescan.
    for name in seed_names_in_order:
        if name not in buckets:
            _ensure(name, seeded_name=name)

    groups: List[ActorGroup] = []
    for key, b in buckets.items():
        seeded_name = b["seeded_name"]
        if seeded_name is not None:
            canonical = seeded_name
            alias_set: Dict[str, None] = {}
            for a in seed_aliases_of.get(canonical, []):
                alias_set[a] = None
            for raw in b["raws"]:
                if normalise_actor_text(raw) == normalise_actor_text(canonical):
                    continue
                if raw not in alias_set:
                    alias_set[raw] = None
            aliases = list(alias_set.keys())
            seeded = True
        else:
            if not b["raws"]:
                continue
            canonical = _pick_canonical(b["raws"])
            aliases = [r for r in b["raws"] if r != canonical]
            seeded = False

        first_seen = None
        if b["first_file"] is not None and b["first_row_ref"] is not None:
            first_seen = f"{b['first_file']} \u2014 {b['first_row_ref']}"

        groups.append(
            ActorGroup(
                canonical=canonical,
                aliases=sorted(aliases, key=str.lower),
                sources=sorted(b["sources"]),
                count=b["count"],
                files=sorted(b["files"]),
                first_seen=first_seen,
                seeded=seeded,
            )
        )

    groups.sort(key=lambda g: (-g.count, g.canonical.lower()))
    return groups


# ---------------------------------------------------------------------------
# Writer — three-sheet xlsx.  Keep styling consistent with writer.py.
# ---------------------------------------------------------------------------


def write_actor_scan(
    result: ActorScanResult,
    output_path: Path,
) -> Path:
    """Write a three-sheet workbook and return the path.

    Sheets:
      * ``Actors`` — Actor / Aliases (+ Count / Files / First seen / Sources).
        Directly readable by ``actors.load_actors_from_xlsx`` (extra columns
        are ignored by that loader).
      * ``Observations`` — every raw sighting with attribution.
      * ``Readme`` — short prose explaining the format and round-trip use.
    """
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Font, PatternFill
    from openpyxl.utils import get_column_letter

    HEADER_FILL = PatternFill("solid", start_color="1F3864")
    HEADER_FONT = Font(name="Arial", bold=True, color="FFFFFF", size=11)
    BODY_FONT = Font(name="Arial", size=10)
    SEED_FILL = PatternFill("solid", start_color="E7F0FB")  # faint blue

    wb = Workbook()

    # --- Actors sheet --- #
    ws = wb.active
    ws.title = "Actors"
    actors_cols = [
        ("Actor", 28),
        ("Aliases", 42),
        ("Count", 8),
        ("Files", 36),
        ("First seen", 34),
        ("Sources", 18),
        ("Seeded", 8),
    ]
    for i, (name, width) in enumerate(actors_cols, start=1):
        cell = ws.cell(row=1, column=i, value=name)
        cell.font = HEADER_FONT
        cell.fill = HEADER_FILL
        cell.alignment = Alignment(horizontal="left", vertical="center")
        ws.column_dimensions[get_column_letter(i)].width = width
    ws.row_dimensions[1].height = 20
    ws.freeze_panes = "A2"

    for row_idx, group in enumerate(result.groups, start=2):
        values = [
            group.canonical,
            ", ".join(group.aliases),
            group.count,
            ", ".join(group.files),
            group.first_seen or "",
            ", ".join(group.sources),
            "yes" if group.seeded else "",
        ]
        for col_idx, value in enumerate(values, start=1):
            cell = ws.cell(row=row_idx, column=col_idx, value=value)
            cell.font = BODY_FONT
            cell.alignment = Alignment(wrap_text=True, vertical="top")
            if group.seeded:
                cell.fill = SEED_FILL

    if result.groups:
        last_col = get_column_letter(len(actors_cols))
        ws.auto_filter.ref = f"A1:{last_col}{len(result.groups) + 1}"

    # --- Observations sheet --- #
    obs_ws = wb.create_sheet("Observations")
    obs_cols = [
        ("Canonical", 28),
        ("Variant", 28),
        ("Source", 12),
        ("File", 28),
        ("Row Ref", 18),
        ("Heading Trail", 40),
    ]
    for i, (name, width) in enumerate(obs_cols, start=1):
        cell = obs_ws.cell(row=1, column=i, value=name)
        cell.font = HEADER_FONT
        cell.fill = HEADER_FILL
        cell.alignment = Alignment(horizontal="left", vertical="center")
        obs_ws.column_dimensions[get_column_letter(i)].width = width
    obs_ws.row_dimensions[1].height = 20
    obs_ws.freeze_panes = "A2"

    # Map every observation to the canonical it ended up in.
    canonical_of: Dict[str, str] = {}
    for group in result.groups:
        ckey = normalise_actor_text(group.canonical)
        canonical_of[ckey] = group.canonical
        for alias in group.aliases:
            canonical_of[normalise_actor_text(alias)] = group.canonical

    for row_idx, obs in enumerate(result.observations, start=2):
        canonical = canonical_of.get(obs.normalised, obs.raw)
        values = [
            canonical,
            obs.raw,
            obs.source,
            obs.file,
            obs.row_ref,
            obs.heading_trail,
        ]
        for col_idx, value in enumerate(values, start=1):
            cell = obs_ws.cell(row=row_idx, column=col_idx, value=value)
            cell.font = BODY_FONT
            cell.alignment = Alignment(wrap_text=True, vertical="top")

    if result.observations:
        last_col = get_column_letter(len(obs_cols))
        obs_ws.auto_filter.ref = f"A1:{last_col}{len(result.observations) + 1}"

    # --- Readme sheet --- #
    readme_ws = wb.create_sheet("Readme")
    readme_ws.column_dimensions["A"].width = 110
    readme_lines = [
        "Actor scan output",
        "",
        "This workbook is a first pass at the list of actors that appear in your",
        "requirements documents.  It is meant to be reviewed and tidied, then",
        "fed back in to a normal extraction run via --actors on the CLI or the",
        "'Actors list' field in the GUI.",
        "",
        "Sheets",
        "------",
        " \u2022 Actors  \u2014 one row per canonical actor.  'Actor' and 'Aliases'",
        "   are the columns load_actors_from_xlsx reads; the remaining columns",
        "   are diagnostic and are ignored on load.",
        " \u2022 Observations \u2014 every raw text span that contributed to a group,",
        "   with its source (primary \u2013 first-column topic; regex \u2013 alias",
        "   match from your seed list; nlp \u2013 spaCy NER hit).",
        "",
        "Round-trip",
        "----------",
        "After you've tidied this file (rename canonicals, move aliases around,",
        "delete rows for false positives), save it and pass the same file to",
        "--actors on your next extraction run.  Seeded rows (highlighted) are",
        "preserved verbatim across re-scans.",
        "",
        "Tip: NLP hits (source='nlp') are best-effort \u2014 entities like 'ISO' or",
        "'the USA' can show up.  Treat them as suggestions, not ground truth.",
    ]
    for i, line in enumerate(readme_lines, start=1):
        cell = readme_ws.cell(row=i, column=1, value=line)
        cell.font = BODY_FONT

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    wb.save(str(output_path))
    return output_path


# ---------------------------------------------------------------------------
# Orchestrator — mirrors extract_from_files's shape so GUI/CLI plumbing is
# trivially parallel.
# ---------------------------------------------------------------------------


def scan_actors_from_files(
    input_paths: Sequence[Path],
    output_path: Path,
    *,
    seed_actors_xlsx: Optional[Path] = None,
    use_nlp: bool = False,
    config_path: Optional[Path] = None,
    progress: Optional[Callable[[str], None]] = None,
    file_progress: Optional[Callable[[int, int, str], None]] = None,
    cancel_check: Optional[Callable[[], bool]] = None,
) -> ActorScanResult:
    """Scan every .docx in ``input_paths`` and write an actors workbook.

    ``seed_actors_xlsx`` optionally seeds the result with an existing
    curated actors list; canonicals and aliases are preserved exactly,
    and any new spellings observed in the docs become additional aliases.

    Honours the same ``cancel_check`` / ``file_progress`` contract as
    :func:`requirements_extractor.extractor.extract_from_files`, so the
    GUI can reuse its plumbing.  A cancelled scan writes nothing and
    raises :class:`ActorScanCancelled`.
    """
    stats = ActorScanStats()
    log = progress or (lambda msg: None)

    seed_entries: List[ActorEntry] = []
    if seed_actors_xlsx is not None:
        try:
            seed_entries = load_actors_from_xlsx(Path(seed_actors_xlsx))
            log(
                f"Loaded {len(seed_entries)} seed actors from "
                f"{Path(seed_actors_xlsx).name}."
            )
        except Exception as e:  # noqa: BLE001
            stats.errors.append(f"Failed to load seed actors: {e}")
            log(f"WARNING: {stats.errors[-1]}")

    resolver = ActorResolver(actors=seed_entries, use_nlp=use_nlp)
    if use_nlp and resolver._nlp is None:  # noqa: SLF001
        stats.errors.append(
            "NLP requested but spaCy (with an English model) is not available. "
            "Install with:  pip install spacy  &&  python -m spacy download en_core_web_sm"
        )
        log(f"WARNING: {stats.errors[-1]}")

    run_config_path: Optional[Path] = Path(config_path) if config_path else None
    if run_config_path is not None:
        try:
            resolve_config(run_config_path=run_config_path, docx_path=None)
            log(f"Loaded run config: {run_config_path.name}")
        except Exception as e:  # noqa: BLE001
            stats.errors.append(f"Failed to load config {run_config_path}: {e}")
            log(f"WARNING: {stats.errors[-1]}")
            run_config_path = None

    observations: List[ActorObservation] = []
    input_list = list(input_paths)
    total_inputs = len(input_list)

    for idx, raw_path in enumerate(input_list, start=1):
        if cancel_check is not None and cancel_check():
            log(f"Cancelled by user after {idx - 1}/{total_inputs} file(s).")
            raise ActorScanCancelled(
                f"Cancelled after {idx - 1}/{total_inputs} file(s)."
            )
        if file_progress is not None:
            file_progress(idx, total_inputs, Path(raw_path).name)

        path = Path(raw_path)
        if not path.exists():
            stats.errors.append(f"File not found: {path}")
            log(f"WARNING: {stats.errors[-1]}")
            continue
        if path.suffix.lower() != ".docx":
            stats.errors.append(f"Skipping non-.docx file: {path.name}")
            log(f"WARNING: {stats.errors[-1]}")
            continue

        try:
            cfg: Config = resolve_config(
                run_config_path=run_config_path, docx_path=path,
            )
        except Exception as e:  # noqa: BLE001
            stats.errors.append(
                f"Failed to load per-doc config for {path.name}: {e}"
            )
            log(f"WARNING: {stats.errors[-1]}")
            cfg = Config.defaults()

        log(f"Scanning {path.name} (config: {cfg.source}) ...")
        try:
            new_obs = _walk_one_doc(path, cfg, resolver)
        except Exception as e:  # noqa: BLE001
            stats.errors.append(f"Error scanning {path.name}: {e}")
            log(f"ERROR: {stats.errors[-1]}")
            continue

        observations.extend(new_obs)
        stats.files_processed += 1
        log(f"  captured {len(new_obs)} observation(s).")

    groups = group_observations(observations, seed_entries=seed_entries)
    stats.observations = len(observations)
    stats.groups = len(groups)

    result = ActorScanResult(
        groups=groups,
        observations=observations,
        stats=stats,
    )

    output_path = Path(output_path)
    write_actor_scan(result, output_path)
    log(
        f"Wrote {len(groups)} actor group(s) from {len(observations)} "
        f"observation(s) to {output_path}."
    )
    result.output_path = output_path
    return result
