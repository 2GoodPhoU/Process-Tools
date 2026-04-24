"""Legacy-format support: convert .doc and .pdf inputs to a form the
main parser can consume.

REVIEW §3.1 extended this session (2026-04-24) to cover:

* **`.doc`** (pre-2007 Word Binary Format).  We shell out to
  LibreOffice's headless converter (``soffice --headless
  --convert-to docx``) to produce a temporary ``.docx`` that the
  existing parser handles natively.
* **`.pdf`** (Portable Document Format).  We use pdfplumber to pull
  any tables that look like the expected actor/content shape into a
  synthetic ``.docx`` with matching 2-column tables, plus carry any
  non-table prose across as preamble paragraphs.  Lossier than the
  Word path — PDFs routinely lose table structure on export — so the
  resulting extraction is best-effort rather than authoritative.

Both paths are gated behind runtime capability checks:

* ``.doc`` needs LibreOffice installed and discoverable on PATH (or
  via a platform-standard install location).  Absent → a ``ValueError``
  with a friendly "install LibreOffice to enable .doc support" message.
* ``.pdf`` needs ``pdfplumber`` installed (listed in
  ``requirements-optional.txt``).  Absent → a ``ValueError`` telling
  the user how to install it.

Neither dependency is required for the default ``.docx`` pipeline —
both are optional.  The :func:`prepare_for_parser` entry point
routes a path to the right conversion, produces a temp ``.docx``,
and returns it along with a cleanup callback.  The caller (typically
``extractor.extract_from_files``) owns the tempdir's lifetime.
"""

from __future__ import annotations

import os
import shutil
import subprocess
import sys
import tempfile
from contextlib import contextmanager
from pathlib import Path
from typing import Iterator, Optional, Tuple


# ---------------------------------------------------------------------------
# LibreOffice discovery
# ---------------------------------------------------------------------------


#: Candidate executable names for LibreOffice headless.  ``soffice`` is
#: the generic name, ``libreoffice`` is the one Debian-based distros
#: expose.  Platform-standard install paths (Windows / macOS) are tried
#: after PATH lookup fails.
_SOFFICE_CANDIDATES: Tuple[str, ...] = ("soffice", "libreoffice")

#: Platform-specific fallbacks — checked only when the PATH lookup
#: misses.  Populated lazily in :func:`find_soffice` so import of this
#: module doesn't touch the filesystem on machines that will never use
#: the .doc path.
_PLATFORM_FALLBACKS = {
    "win32": (
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
    ),
    "darwin": (
        "/Applications/LibreOffice.app/Contents/MacOS/soffice",
    ),
    # Linux ships LibreOffice via package managers, which put the
    # binaries on PATH.  No fallbacks needed.
}


def find_soffice() -> Optional[str]:
    """Return the full path to a working LibreOffice binary, or None.

    Checks PATH first, then the platform-specific fallbacks for
    Windows and macOS.  Cached via :func:`functools.lru_cache` would
    be overkill here — look-ups are rare (once per .doc input at
    most) and the answer can legitimately change mid-session (user
    installs LibreOffice while the GUI is running).
    """
    for name in _SOFFICE_CANDIDATES:
        resolved = shutil.which(name)
        if resolved:
            return resolved
    for path in _PLATFORM_FALLBACKS.get(sys.platform, ()):
        if os.path.exists(path):
            return path
    return None


def has_soffice() -> bool:
    """Lightweight Boolean wrapper — "can we do .doc conversion?"."""
    return find_soffice() is not None


class LibreOfficeUnavailable(RuntimeError):
    """Raised when a .doc input is given but LibreOffice is not installed."""


def _missing_libreoffice_message() -> str:
    """Build the platform-appropriate install-instruction string."""
    base = (
        "LibreOffice is required to read legacy .doc files.  Install it "
        "once per machine and re-run:\n"
    )
    if sys.platform.startswith("win"):
        return base + (
            "  Windows:  https://www.libreoffice.org/download/download/\n"
            "  After install, confirm ``soffice.exe`` exists in "
            "``C:\\Program Files\\LibreOffice\\program\\``."
        )
    if sys.platform == "darwin":
        return base + (
            "  macOS:    brew install --cask libreoffice\n"
            "  (Or download the .dmg from https://www.libreoffice.org/.)"
        )
    return base + (
        "  Debian/Ubuntu:  sudo apt install libreoffice\n"
        "  Fedora/RHEL:    sudo dnf install libreoffice\n"
        "  Arch:           sudo pacman -S libreoffice-fresh"
    )


# ---------------------------------------------------------------------------
# .doc conversion
# ---------------------------------------------------------------------------


def convert_doc_to_docx(
    doc_path: Path,
    output_dir: Path,
    *,
    timeout_s: int = 120,
) -> Path:
    """Convert ``doc_path`` (``.doc``) to ``.docx`` via LibreOffice.

    The converted file lands in ``output_dir`` with the same stem as
    the input and a ``.docx`` extension.  Returns the path of the
    produced file.

    Raises:
      * :class:`LibreOfficeUnavailable` if no soffice binary can be
        found — message points the user at install instructions.
      * :class:`RuntimeError` if the conversion process runs but the
        expected output file isn't written (LibreOffice silently
        failing on a corrupt .doc is the usual cause).
      * :class:`subprocess.TimeoutExpired` if the conversion takes
        longer than ``timeout_s`` seconds — default 120 is generous
        for normal docs; corrupt inputs sometimes hang otherwise.
    """
    soffice = find_soffice()
    if soffice is None:
        raise LibreOfficeUnavailable(_missing_libreoffice_message())

    doc_path = Path(doc_path)
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    expected_out = output_dir / (doc_path.stem + ".docx")

    # LibreOffice's CLI quirk: ``--outdir`` decides where the file
    # lands.  It always names the output ``<stem>.docx`` regardless of
    # what the input was called (``.doc``, ``.rtf``, etc.), which
    # matches what we want.  We route the subprocess's stdout/stderr
    # into pipes so the caller's terminal stays clean and so we can
    # surface meaningful diagnostics if the conversion fails.
    cmd = [
        soffice,
        "--headless",
        "--convert-to", "docx",
        "--outdir", str(output_dir),
        str(doc_path),
    ]
    result = subprocess.run(
        cmd,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        timeout=timeout_s,
        check=False,
    )
    if result.returncode != 0 or not expected_out.exists():
        stderr = (result.stderr or b"").decode("utf-8", errors="replace").strip()
        stdout = (result.stdout or b"").decode("utf-8", errors="replace").strip()
        detail = stderr or stdout or f"return code {result.returncode}"
        raise RuntimeError(
            f"LibreOffice failed to convert {doc_path.name} to .docx: "
            f"{detail}"
        )
    return expected_out


# ---------------------------------------------------------------------------
# .pdf conversion (pdfplumber-backed, best-effort)
# ---------------------------------------------------------------------------


class PdfSupportUnavailable(RuntimeError):
    """Raised when a .pdf input is given but pdfplumber isn't installed."""


def _missing_pdfplumber_message() -> str:
    return (
        "pdfplumber is required to read .pdf files.  Install with:\n"
        "    pip install pdfplumber\n"
        "or add ``pdfplumber`` to your requirements-optional.txt and "
        "re-run ``pip install -r requirements-optional.txt``."
    )


def convert_pdf_to_docx(
    pdf_path: Path,
    output_dir: Path,
) -> Path:
    """Convert ``pdf_path`` to a synthetic ``.docx`` the parser can consume.

    Strategy:
      * Open the PDF with pdfplumber.
      * For every detected table on every page, emit it as a native
        Word table in the output ``.docx``.  No attempt is made to
        guess which columns are actor vs. content — the extractor's
        per-doc ``.reqx.yaml`` mechanism is the right place for that.
        A 2-column table (the default configuration) falls through
        the normal pipeline with zero additional config.
      * For pages **with no tables**, emit the page text as plain
        paragraphs so the preamble path picks them up.
      * For pages **with tables**, skip the page-text extraction
        entirely — the table content already carried the structured
        form, and extracting the text again would produce duplicate
        requirements in the output (the same sentence would appear
        once as a table-row requirement and once as a preamble
        requirement).  When the author structured content as a table,
        trust that structure.  Users who want the surrounding prose
        captured anyway can convert the PDF to ``.docx`` via another
        tool (LibreOffice, Acrobat export) and run the extractor on
        that.

    This is lossier than native ``.docx`` in every way — page
    headers/footers, reading order across columns, and any ASCII art
    that pdfplumber didn't identify as a table will all be wrong.
    It's useful for a "we only have the PDF" scenario; if the source
    Word document is available, prefer that.

    Raises:
      * :class:`PdfSupportUnavailable` if pdfplumber isn't importable.
      * :class:`RuntimeError` if the PDF itself cannot be opened (bad
        file, password-protected, etc.).
    """
    try:
        import pdfplumber  # type: ignore
    except ImportError as exc:
        raise PdfSupportUnavailable(_missing_pdfplumber_message()) from exc

    from docx import Document as DocxDocument

    pdf_path = Path(pdf_path)
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    out_path = output_dir / (pdf_path.stem + ".docx")

    doc = DocxDocument()

    try:
        with pdfplumber.open(str(pdf_path)) as pdf:
            for page_num, page in enumerate(pdf.pages, start=1):
                # Extract tables first — they're the most structured
                # part of any spec PDF and the part we most want the
                # parser to see.
                tables = page.extract_tables() or []
                for table in tables:
                    # pdfplumber returns tables as list-of-lists, each
                    # row a list of cell strings (may be None for
                    # empty cells).  Normalise to empty strings so
                    # python-docx doesn't emit literal "None".
                    clean = [
                        [("" if c is None else str(c).strip()) for c in row]
                        for row in table or []
                        if row is not None
                    ]
                    if not clean:
                        continue
                    n_cols = max(len(r) for r in clean)
                    word_table = doc.add_table(rows=len(clean), cols=n_cols)
                    word_table.style = "Table Grid"
                    for r_idx, row in enumerate(clean):
                        for c_idx, value in enumerate(row):
                            cell = word_table.rows[r_idx].cells[c_idx]
                            cell.paragraphs[0].text = value

                # Page-text extraction only runs when we found no
                # tables on this page.  Rationale: pdfplumber's
                # extract_text() pulls EVERY glyph on the page —
                # including the ones that were already captured by
                # the table extractor above.  Emitting both would
                # double-count every requirement (see the note in the
                # docstring).  For a spec page where the author used a
                # table to carry the requirements, the table IS the
                # intended structure; surrounding page headers and
                # footers are usually not useful anyway.
                #
                # For pages without tables (intro prose, glossary,
                # references), we still emit the text as paragraphs
                # so the preamble pipeline picks up anything
                # shall-matched.
                if tables:
                    # Light marker so a reviewer grepping "-- PDF
                    # page N --" can still tell which page produced
                    # which tables.
                    doc.add_paragraph(
                        f"-- PDF page {page_num} (tables only) --"
                    )
                    continue
                text = page.extract_text() or ""
                if text.strip():
                    doc.add_paragraph(f"-- PDF page {page_num} --")
                    for line in text.splitlines():
                        line = line.strip()
                        if line:
                            doc.add_paragraph(line)
    except PdfSupportUnavailable:
        raise
    except Exception as exc:  # noqa: BLE001 — best-effort PDF read
        raise RuntimeError(
            f"Failed to read PDF {pdf_path.name}: {exc}"
        ) from exc

    doc.save(str(out_path))
    return out_path


# ---------------------------------------------------------------------------
# Unified entry point
# ---------------------------------------------------------------------------


@contextmanager
def prepare_for_parser(path: Path) -> Iterator[Path]:
    """Yield a ``.docx`` path suitable for the main parser.

    Behaviour by extension:
      * ``.docx``  - yields the input path unchanged.  No conversion,
        no tempdir.
      * ``.doc``   - yields a converted temp ``.docx``.  Cleans up the
        tempdir on exit.
      * ``.pdf``   - yields a synthetic temp ``.docx`` built from
        pdfplumber-extracted tables and prose.  Cleans up on exit.

    Used as::

        with prepare_for_parser(path) as ready:
            events = parse_docx_events(ready, ...)

    Any extension other than the three above raises ``ValueError`` -
    the caller should surface it as a user-facing warning and skip
    the file.
    """
    suffix = path.suffix.lower()
    if suffix == ".docx":
        # Fast path - no tempdir allocation for the common case.
        yield Path(path)
        return
    if suffix not in {".doc", ".pdf"}:
        raise ValueError(
            f"Unsupported input extension: {suffix!r}.  Expected "
            f"one of .docx / .doc / .pdf."
        )
    with tempfile.TemporaryDirectory(prefix="docx_extractor_legacy_") as td:
        td_path = Path(td)
        if suffix == ".doc":
            converted = convert_doc_to_docx(Path(path), td_path)
        else:  # ".pdf"
            converted = convert_pdf_to_docx(Path(path), td_path)
        yield converted
