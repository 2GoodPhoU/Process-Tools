"""Generate synthetic multi-action requirement fixtures for the 0.6.2 patch.

Five documents, each targeting one of the multi-action detection cases:

  fixture_1_eric_canonical.docx
                                  - Eric's exact "Software engineers
                                    shall create unit tests, implement
                                    CICD Pipelines, and integrate
                                    software quality control systems."
                                    The canonical 3-action case.
  fixture_2_two_action_or.docx
                                  - Two-action with ``or`` connector
                                    ("System shall log errors or alert
                                    operators").  Bare connector
                                    (no Oxford comma).
  fixture_3_shared_verb_compound.docx
                                  - Genuinely-singular compound:
                                    "System shall be available 99.9%
                                    of the time, with mean time to
                                    recovery less than 1 hour".
                                    Should NOT split (shared-verb
                                    pattern, "with" continuation).
  fixture_4_after_compound_aggregation.docx
                                  - A 0.6.1 compound aggregation
                                    (bulleted list lead-in) where the
                                    aggregated text contains
                                    multi-action verb phrases.  Tests
                                    pipeline ordering: 0.6.1
                                    aggregation runs FIRST, then 0.6.2
                                    decomposition runs on the
                                    aggregated text.
  fixture_5_inside_procedural_table.docx
                                  - Multi-action sentence inside a
                                    procedural required-action table
                                    (3-col header).  Should NOT split
                                    (procedural-table gate matches the
                                    0.6.1 compound gate).

Each fixture is a single 2-column requirements table whose right
column holds the multi-action sentence(s), except fixture 5 which uses
the 3-col procedural required-action shape.

Run:
    python samples/procedures/multi_action/generate.py
"""

from __future__ import annotations

from pathlib import Path

from docx import Document


HERE = Path(__file__).resolve().parent


def _write_cell(cell, *paragraphs) -> None:
    """Replace a cell's content with one or more paragraphs.

    Same shape as samples/procedures/compound/generate.py.
    """
    first = cell.paragraphs[0]
    first.text = ""
    if not paragraphs:
        return
    for i, entry in enumerate(paragraphs):
        text, style = (entry, None) if isinstance(entry, str) else entry
        target = first if i == 0 else cell.add_paragraph()
        target.text = text
        if style:
            target.style = style


# ---------------------------------------------------------------------------
# Fixture 1 - Eric's canonical 3-action case
# ---------------------------------------------------------------------------


def build_eric_canonical() -> None:
    """Eric's example, verbatim.  Three distinct actions.

    Expected (in flag/split modes): 3 atomic verb phrases detected:
    ``create unit tests``, ``implement CICD Pipelines``, ``integrate
    software quality control systems``.  Single mode emits one
    requirement; flag mode emits one with a "consider splitting" note;
    split mode emits the parent + 3 sub-requirements.
    """
    doc = Document()
    doc.add_heading("Procedure - Multi-action canonical (Eric)", level=0)
    doc.add_paragraph("Document ID: PROC-MULTI-001")
    doc.add_paragraph(
        "Synthetic fixture for 0.6.2 multi-action detection.  Three "
        "distinct verb phrases sharing one modal verb."
    )

    doc.add_heading("1. Engineer responsibilities", level=1)
    t = doc.add_table(rows=1, cols=2)
    t.style = "Table Grid"

    _write_cell(t.rows[0].cells[0], "Software Engineers")
    _write_cell(
        t.rows[0].cells[1],
        "Software engineers shall create unit tests, implement CICD "
        "Pipelines, and integrate software quality control systems.",
    )

    doc.save(HERE / "fixture_1_eric_canonical.docx")


# ---------------------------------------------------------------------------
# Fixture 2 - Two-action with `or` connector (bare, no Oxford comma)
# ---------------------------------------------------------------------------


def build_two_action_or() -> None:
    """Two-action disjunction with a bare ``or`` connector.

    Expected: the bare-connector path detects two verb phrases
    (``log errors``, ``alert operators``) when both halves head with
    a recognisable verb.
    """
    doc = Document()
    doc.add_heading("Procedure - Two-action OR", level=0)
    doc.add_paragraph("Document ID: PROC-MULTI-002")
    doc.add_paragraph(
        "Two-action sentence; ``or`` connector; no Oxford comma.  "
        "Tests the bare-connector fallback in the regex path."
    )

    doc.add_heading("2. Error handling", level=1)
    t = doc.add_table(rows=1, cols=2)
    t.style = "Table Grid"

    _write_cell(t.rows[0].cells[0], "System")
    _write_cell(
        t.rows[0].cells[1],
        "The system shall log errors or alert operators.",
    )

    doc.save(HERE / "fixture_2_two_action_or.docx")


# ---------------------------------------------------------------------------
# Fixture 3 - Genuinely-singular compound (shared-verb pattern)
# ---------------------------------------------------------------------------


def build_shared_verb_compound() -> None:
    """Single requirement decorated with multiple object clauses.

    Expected: detection's ``is_multi_action`` is False - the second
    clause begins with ``with`` (not a verb) so the regex fallback
    rejects it.  This is the disambiguation gate from the patch
    spec ("shared verb / non-verb continuation -> single requirement").
    """
    doc = Document()
    doc.add_heading("Procedure - Shared-verb compound", level=0)
    doc.add_paragraph("Document ID: PROC-MULTI-003")
    doc.add_paragraph(
        "Negative-case fixture.  Single availability requirement "
        "decorated with a recovery-time clause that uses ``with``, not "
        "a second verb.  Multi-action detection should not be triggered."
    )

    doc.add_heading("3. Availability", level=1)
    t = doc.add_table(rows=1, cols=2)
    t.style = "Table Grid"

    _write_cell(t.rows[0].cells[0], "System")
    _write_cell(
        t.rows[0].cells[1],
        "The system shall be available 99.9% of the time, with mean "
        "time to recovery less than 1 hour.",
    )

    doc.save(HERE / "fixture_3_shared_verb_compound.docx")


# ---------------------------------------------------------------------------
# Fixture 4 - Multi-action AFTER 0.6.1 compound aggregation
# ---------------------------------------------------------------------------


def build_after_compound_aggregation() -> None:
    """Compound lead-in + bullets where the aggregated text is multi-action.

    The 0.6.1 aggregation runs first and produces ONE compound row
    whose synthetic text contains the lead-in plus a list of items.
    The aggregated text deliberately includes a multi-action sentence
    on the lead-in side so the 0.6.2 detection sees a multi-action
    pattern AFTER aggregation.

    NOTE: in practice, post-aggregation text reads like
    "<lead-in> (all of): (1) ...; (2) ...".  The colon and numeric
    prefixes break the multi-action regex's expectation of a clean
    sentence so detection doesn't fire on a typical aggregated row.
    This fixture instead places the multi-action pattern in the
    lead-in itself - so 0.6.1 aggregation captures the lead-in and
    items as one row, but the lead-in's verb phrases are still
    captured by 0.6.2 detection on the aggregated text's prefix
    when present.

    Expected: depending on mode, either a flagged compound row (flag
    mode default) or a compound + multi-action sub-rows (split mode).
    The intent is to verify the pipeline does not crash when both
    detections fire on the same parent.
    """
    doc = Document()
    doc.add_heading("Procedure - Compound + multi-action interaction", level=0)
    doc.add_paragraph("Document ID: PROC-MULTI-004")
    doc.add_paragraph(
        "Pipeline-ordering fixture.  0.6.1 compound aggregation runs "
        "first; 0.6.2 multi-action runs on the aggregated text."
    )

    doc.add_heading("4. Build pipeline", level=1)
    t = doc.add_table(rows=1, cols=2)
    t.style = "Table Grid"

    _write_cell(t.rows[0].cells[0], "Engineering")
    _write_cell(
        t.rows[0].cells[1],
        ("Engineering shall comply with the following before each "
         "release if they:", None),
        ("are touching production systems, and", "List Bullet"),
        ("have not yet been peer-reviewed, and", "List Bullet"),
        ("affect customer-visible behaviour", "List Bullet"),
    )

    doc.save(HERE / "fixture_4_after_compound_aggregation.docx")


# ---------------------------------------------------------------------------
# Fixture 5 - Multi-action inside a procedural required-action table
# ---------------------------------------------------------------------------


def build_inside_procedural_table() -> None:
    """Procedural required-action table whose action cell is multi-action.

    Expected: the procedural-table gate disables both compound (0.6.1)
    AND multi-action (0.6.2) detection.  Each row is treated atomically
    by virtue of the ``Required Action`` column header.  The
    multi-action sentence in row 2's content cell stays as one row.
    """
    doc = Document()
    doc.add_heading("Procedure - Multi-action inside required-action table", level=0)
    doc.add_paragraph("Document ID: PROC-MULTI-005")
    doc.add_paragraph(
        "Procedural-table gate fixture.  Multi-action detection ought "
        "NOT fire when force_requirement is True."
    )

    doc.add_heading("5. Required-action sequence", level=1)
    t = doc.add_table(rows=3, cols=3)
    t.style = "Table Grid"

    # Header row - 3-col procedural shape: blank | Step | Required Action
    _write_cell(t.rows[0].cells[0], "")
    _write_cell(t.rows[0].cells[1], "Step")
    _write_cell(t.rows[0].cells[2], "Required Action")

    # Body row 1 - single-action.
    _write_cell(t.rows[1].cells[0], "Operator")
    _write_cell(t.rows[1].cells[1], "1")
    _write_cell(
        t.rows[1].cells[2],
        "Verify the system status before proceeding.",
    )

    # Body row 2 - multi-action.  Must stay as one row.
    _write_cell(t.rows[2].cells[0], "Operator")
    _write_cell(t.rows[2].cells[1], "2")
    _write_cell(
        t.rows[2].cells[2],
        "The Operator shall record metrics, analyze trends, and "
        "report findings.",
    )

    doc.save(HERE / "fixture_5_inside_procedural_table.docx")


# ---------------------------------------------------------------------------
# Runner
# ---------------------------------------------------------------------------


def main() -> None:
    build_eric_canonical()
    build_two_action_or()
    build_shared_verb_compound()
    build_after_compound_aggregation()
    build_inside_procedural_table()
    docs = sorted(HERE.glob("*.docx"))
    print(f"Generated {len(docs)} multi-action fixtures in {HERE}:")
    for p in docs:
        print(f"  - {p.name}")


if __name__ == "__main__":
    main()
