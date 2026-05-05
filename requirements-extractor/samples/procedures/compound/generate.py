"""Generate synthetic compound-requirement fixtures for the 0.6.1 patch.

Five documents, each targeting one of the compound-detection cases:

  fixture_1_and_list.docx        — Eric's "Programs shall comply..."
                                   example: bulleted list of AND-joined
                                   conditions following a modal+colon
                                   lead-in.  The canonical case.
  fixture_2_or_list.docx         — Same shape, but items are joined by
                                   ``or`` (any-of disjunction).
  fixture_3_mixed_markers.docx   — Numbered list (``1.``, ``2.``, etc.)
                                   instead of bullets — exercises the
                                   plain-text-marker fallback.  Also
                                   throws in a hand-typed dash item.
  fixture_4_where_definitions.docx
                                  — ``where:`` lead-in introduces a
                                    glossary block.  Compound detection
                                    must NOT fire; the existing
                                    per-paragraph walk handles each
                                    definition separately.
  fixture_5_unless_exclusion.docx
                                  — ``unless`` connector, polarity
                                    flipped: items are conditions that
                                    EXCLUDE the obligation.

Each fixture is a single 2-column requirements table whose right column
holds the lead-in + list.  This shape matches Eric's 2026-04-23 field
notes — it's how the real procedure documents are authored.

Run:
    python samples/procedures/compound/generate.py
"""

from __future__ import annotations

from pathlib import Path

from docx import Document


HERE = Path(__file__).resolve().parent


def _write_cell(cell, *paragraphs) -> None:
    """Replace a cell's content with one or more paragraphs.

    Same shape as ``samples/procedures/generate.py``: each entry is
    either a plain ``str`` (unstyled body paragraph) or a
    ``(text, style_name)`` tuple that applies a Word style — used for
    bullet / numbered list items inside a cell.
    """
    first = cell.paragraphs[0]
    first.text = ""
    if not paragraphs:
        return
    for i, entry in enumerate(paragraphs):
        text, style = (entry, None) if isinstance(entry, str) else entry
        target = first if i == 0 else cell.add_paragraph()
        target.text = text
        if style:
            target.style = style


# ---------------------------------------------------------------------------
# Fixture 1 — "Programs shall comply..." (AND list, bullets)
# ---------------------------------------------------------------------------


def build_and_list() -> None:
    """Eric's example, almost verbatim.  Bullets joined by ``and``.

    Expected: ONE compound requirement covering the lead-in plus four
    AND-joined conditions.  In 0.6.0 this would have produced zero
    requirements (lead-in ends with ``:`` so detector treated it as
    fragmentary; bullets had no modal verb).
    """
    doc = Document()
    doc.add_heading("Procedure — Compound AND list", level=0)
    doc.add_paragraph("Document ID: PROC-COMP-001")
    doc.add_paragraph(
        "Synthetic fixture for 0.6.1 compound-requirement detection.  "
        "AND-joined conditions following a modal+colon lead-in."
    )

    doc.add_heading("1. Eligibility", level=1)
    t = doc.add_table(rows=1, cols=2)
    t.style = "Table Grid"

    _write_cell(t.rows[0].cells[0], "Program Office")
    _write_cell(
        t.rows[0].cells[1],
        ("Programs shall comply with this document if they:", None),
        ("Domestic US-Based program, and", "List Bullet"),
        ("exceed 12 months in duration, and", "List Bullet"),
        ("Pre-preliminary design review <MRL 5 overall maturity, and",
         "List Bullet"),
        ("design and/produce hardware content", "List Bullet"),
    )

    doc.save(HERE / "fixture_1_and_list.docx")


# ---------------------------------------------------------------------------
# Fixture 2 — OR list (any-of disjunction)
# ---------------------------------------------------------------------------


def build_or_list() -> None:
    """Disjunctive list: any one of the conditions triggers the obligation.

    Expected: ONE compound requirement; connector inferred as ``or``;
    the synthetic compound text reads "...(any of): (1) ...; (2) ...;
    (3) ...".
    """
    doc = Document()
    doc.add_heading("Procedure — Compound OR list", level=0)
    doc.add_paragraph("Document ID: PROC-COMP-002")
    doc.add_paragraph(
        "Synthetic fixture: any-of disjunction.  Items end with `or` "
        "to signal the connector."
    )

    doc.add_heading("2. Notification triggers", level=1)
    t = doc.add_table(rows=1, cols=2)
    t.style = "Table Grid"

    _write_cell(t.rows[0].cells[0], "Operations")
    _write_cell(
        t.rows[0].cells[1],
        ("Operations must notify the on-call supervisor if any of the following occur:",
         None),
        ("subsystem latency exceeds 200 ms, or", "List Bullet"),
        ("error budget burn-rate doubles within a 5-minute window, or",
         "List Bullet"),
        ("synthetic monitor returns a non-2xx response", "List Bullet"),
    )

    doc.save(HERE / "fixture_2_or_list.docx")


# ---------------------------------------------------------------------------
# Fixture 3 — Mixed-marker numbered + hand-typed dash list
# ---------------------------------------------------------------------------


def build_mixed_markers() -> None:
    """Numbered list (``List Number`` style) + hand-typed dash item.

    Exercises both Word's ``numPr`` signal (the styled items) and the
    plain-text marker fallback (the hand-typed dash item).  Expected:
    one compound spanning all four items, joined by ``and``.
    """
    doc = Document()
    doc.add_heading("Procedure — Mixed list markers", level=0)
    doc.add_paragraph("Document ID: PROC-COMP-003")
    doc.add_paragraph(
        "Numbered list with a hand-typed dash item appended.  Both "
        "marker styles must be claimed by the same compound group."
    )

    doc.add_heading("3. Pre-flight checks", level=1)
    t = doc.add_table(rows=1, cols=2)
    t.style = "Table Grid"

    _write_cell(t.rows[0].cells[0], "Pilot")
    _write_cell(
        t.rows[0].cells[1],
        ("The Pilot shall complete the following pre-flight checks before takeoff:",
         None),
        ("verify fuel reserves match the flight plan, and",
         "List Number"),
        ("confirm all control surfaces respond freely, and",
         "List Number"),
        ("test the radio handshake with ground control, and",
         "List Number"),
        ("- record the takeoff weight in the journey log", None),
    )

    doc.save(HERE / "fixture_3_mixed_markers.docx")


# ---------------------------------------------------------------------------
# Fixture 4 — `where:` definition list (must NOT aggregate)
# ---------------------------------------------------------------------------


def build_where_definitions() -> None:
    """A modal+colon lead-in whose last word is ``where``.

    This shape introduces a glossary / definitions block; the compound
    detector must NOT fire.  Expected: the lead-in passes through to
    the legacy walker (which captures it as a regular requirement
    because it has a modal); the bullets fall through and are emitted
    as bullet rows by the legacy path.

    Counts as a NEGATIVE test: 0.6.1 must NOT aggregate this block,
    even though it superficially resembles fixture 1.
    """
    doc = Document()
    doc.add_heading("Procedure — Where-clause definitions", level=0)
    doc.add_paragraph("Document ID: PROC-COMP-004")
    doc.add_paragraph(
        "Negative-case fixture.  The lead-in's last word is `where`, "
        "which the compound detector must treat as a definition block "
        "rather than an aggregated condition list."
    )

    doc.add_heading("4. Notation", level=1)
    t = doc.add_table(rows=1, cols=2)
    t.style = "Table Grid"

    _write_cell(t.rows[0].cells[0], "System")
    _write_cell(
        t.rows[0].cells[1],
        ("The system shall enforce the throughput envelope T(x) where:",
         None),
        ("T is the measured throughput in requests per second",
         "List Bullet"),
        ("x is the concurrent-session count at sample time",
         "List Bullet"),
        ("the envelope is published in the operations runbook",
         "List Bullet"),
    )

    doc.save(HERE / "fixture_4_where_definitions.docx")


# ---------------------------------------------------------------------------
# Fixture 5 — `unless:` exclusion list
# ---------------------------------------------------------------------------


def build_unless_exclusion() -> None:
    """Exclusion list: items end with ``unless``, flipping polarity.

    Expected: ONE compound requirement; connector inferred as
    ``unless``; synthetic compound text reads "...(unless any of): ...".
    The connector word MUST be preserved in the rendered compound so
    reviewers see at a glance that these are exclusion conditions, not
    conjunctions.
    """
    doc = Document()
    doc.add_heading("Procedure — Unless-exclusion list", level=0)
    doc.add_paragraph("Document ID: PROC-COMP-005")
    doc.add_paragraph(
        "Exclusion list.  The trailing `unless` words mark the items "
        "as conditions that suspend the lead-in obligation."
    )

    doc.add_heading("5. Maintenance window", level=1)
    t = doc.add_table(rows=1, cols=2)
    t.style = "Table Grid"

    _write_cell(t.rows[0].cells[0], "Operations")
    _write_cell(
        t.rows[0].cells[1],
        ("Operations shall apply the patch within 24 hours of release "
         "unless:", None),
        ("the change touches a regulated subsystem, unless",
         "List Bullet"),
        ("the maintenance window has been deferred by the Change Board, "
         "unless", "List Bullet"),
        ("an active incident covers the affected service",
         "List Bullet"),
    )

    doc.save(HERE / "fixture_5_unless_exclusion.docx")


# ---------------------------------------------------------------------------
# Runner
# ---------------------------------------------------------------------------


def main() -> None:
    build_and_list()
    build_or_list()
    build_mixed_markers()
    build_where_definitions()
    build_unless_exclusion()
    docs = sorted(HERE.glob("*.docx"))
    print(f"Generated {len(docs)} compound fixtures in {HERE}:")
    for p in docs:
        print(f"  - {p.name}")


if __name__ == "__main__":
    main()
