"""Generate synthetic procedure documents exercising actor-ID failure modes.

FIELD_NOTES §4. The real procedure documents the tool is run against
are controlled and can't be committed. These hand-authored synthetic
analogues reproduce the *structural* patterns that drive actor-ID
errors without containing any controlled content — industry-generic
shapes only.

Five documents, each targeting a specific failure mode:

  simple_two_actors.docx       — baseline: 2 named actors, clean prose.
                                 Pins the "it still works on easy mode"
                                 regression.
  ambiguous_roles.docx         — nested procedure where the same actor
                                 appears under several spellings
                                 ("Ops Engineer", "Operations Engineer",
                                 "the engineer") and multiple roles blur
                                 into one step.
  implicit_system_actor.docx   — every step attributed to "the system"
                                 or omitted entirely. Primary-column is
                                 empty or generic. Tests whether the
                                 resolver surfaces a reasonable actor
                                 when none is named.
  passive_voice.docx           — heavy passive-voice prose ("shall be
                                 performed by") that hides the agent
                                 inside a by-phrase. NLP should still
                                 catch most; regex/heuristic will miss.
  parallel_flows.docx          — multi-step parallel flows where two
                                 actors progress independently in
                                 interleaved steps. Tests row-ordering
                                 and actor continuity.

Run from the project root:
    python samples/procedures/generate.py

Output .docx files are written next to this script so the generator
and its artefacts live together — same pattern as samples/edge_cases/.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.table import _Cell


HERE = Path(__file__).resolve().parent


# ---------------------------------------------------------------------------
# Small helpers — mirror samples/edge_cases/generate.py's conventions so
# anyone who's read one generator can read the other.
# ---------------------------------------------------------------------------


def _write_cell(cell: _Cell, *paragraphs) -> None:
    """Replace a cell's content with one or more paragraphs.

    Each entry in ``paragraphs`` is either a plain ``str`` (treated as
    an unstyled body paragraph) or a ``(text, style_name)`` tuple which
    applies the named Word style — useful for bulleted / numbered list
    items inside a cell.  Accepts the same shape as
    ``samples/edge_cases/generate.py`` so readers of either generator
    don't have to re-learn conventions.
    """
    first = cell.paragraphs[0]
    first.text = ""
    if not paragraphs:
        return
    for i, entry in enumerate(paragraphs):
        text, style = (entry, None) if isinstance(entry, str) else entry
        target = first if i == 0 else cell.add_paragraph()
        target.text = text
        if style:
            target.style = style


def _add_heading(doc: Document, text: str, level: int) -> None:
    doc.add_heading(text, level=level)


# ---------------------------------------------------------------------------
# Sample 1 — Simple two-actor procedure (baseline)
# ---------------------------------------------------------------------------


def build_simple_two_actors() -> None:
    """Two actors, clean prose, one requirement per row.

    The regression test: if this one stops working cleanly, something
    fundamental regressed. Expect 4 requirements (2 per actor), all
    Hard (shall / must).
    """
    doc = Document()
    doc.add_heading("Procedure \u2014 Daily System Check", level=0)
    doc.add_paragraph("Document ID: PROC-BASE-001")
    doc.add_paragraph(
        "Describes the shift-handover checks performed jointly by the "
        "Operator and the Supervisor."
    )

    _add_heading(doc, "1. Shift Start", level=1)
    t = doc.add_table(rows=4, cols=2)
    t.style = "Table Grid"

    _write_cell(t.rows[0].cells[0], "Operator")
    _write_cell(
        t.rows[0].cells[1],
        "The Operator shall log in to the control console and confirm "
        "the previous shift's handover notes.",
    )

    _write_cell(t.rows[1].cells[0], "Supervisor")
    _write_cell(
        t.rows[1].cells[1],
        "The Supervisor must review the overnight alarm log before "
        "releasing the Operator to normal duties.",
    )

    _write_cell(t.rows[2].cells[0], "Operator")
    _write_cell(
        t.rows[2].cells[1],
        "The Operator shall verify that all subsystem status indicators "
        "read nominal.",
    )

    _write_cell(t.rows[3].cells[0], "Supervisor")
    _write_cell(
        t.rows[3].cells[1],
        "The Supervisor must countersign the handover log once the "
        "checks are complete.",
    )

    doc.save(HERE / "simple_two_actors.docx")


# ---------------------------------------------------------------------------
# Sample 2 — Ambiguous actor spellings
# ---------------------------------------------------------------------------


def build_ambiguous_roles() -> None:
    """Same actor appears under several spellings; nested procedure.

    Exercises the resolver's normalisation + grouping rules. The
    aliasing is deliberately messy so the test can pin down which
    normalisations are working and which aren't.
    """
    doc = Document()
    doc.add_heading("Procedure \u2014 Incident Escalation", level=0)
    doc.add_paragraph("Document ID: PROC-AMBIG-002")
    doc.add_paragraph(
        "Escalation path for production incidents. The same roles are "
        "referenced under multiple spellings across the tables."
    )

    _add_heading(doc, "2. Triage", level=1)

    outer = doc.add_table(rows=2, cols=2)
    outer.style = "Table Grid"

    _write_cell(outer.rows[0].cells[0], "Ops Engineer")
    r0c1 = outer.rows[0].cells[1]
    _write_cell(
        r0c1,
        "The Ops Engineer shall open a triage channel within 5 minutes "
        "of page acknowledgement.",
    )
    nested = r0c1.add_table(rows=3, cols=2)
    nested.style = "Table Grid"

    _write_cell(nested.rows[0].cells[0], "Operations Engineer")
    _write_cell(
        nested.rows[0].cells[1],
        "The Operations Engineer must capture the initial alert payload "
        "and paste it into the channel.",
    )
    _write_cell(nested.rows[1].cells[0], "the engineer")
    _write_cell(
        nested.rows[1].cells[1],
        "The engineer shall page the on-call Platform Lead if the "
        "incident is not acknowledged within 10 minutes.",
    )
    _write_cell(nested.rows[2].cells[0], "Platform Lead")
    _write_cell(
        nested.rows[2].cells[1],
        "The Platform Lead must join the triage channel and assume "
        "incident command.",
    )

    _write_cell(outer.rows[1].cells[0], "Platform-Lead")
    _write_cell(
        outer.rows[1].cells[1],
        "The Platform-Lead shall designate a scribe and a communicator "
        "within the first 15 minutes of the incident.",
    )

    doc.save(HERE / "ambiguous_roles.docx")


# ---------------------------------------------------------------------------
# Sample 3 — Implicit / unnamed actor ("the system")
# ---------------------------------------------------------------------------


def build_implicit_system_actor() -> None:
    """Steps attributed to 'the system' or left unattributed.

    Primary-column cells are either generic ("System") or blank, which
    is the documentary reality for auto-generated or loosely-authored
    procedures. Tests whether the resolver returns anything sensible.
    """
    doc = Document()
    doc.add_heading("Procedure \u2014 Automated Backup Run", level=0)
    doc.add_paragraph("Document ID: PROC-IMPLICIT-003")
    doc.add_paragraph(
        "Describes the nightly backup sequence. Steps are largely "
        "automated; explicit actor naming is minimal."
    )

    _add_heading(doc, "3. Sequence", level=1)
    t = doc.add_table(rows=5, cols=2)
    t.style = "Table Grid"

    _write_cell(t.rows[0].cells[0], "System")
    _write_cell(
        t.rows[0].cells[1],
        "At 02:00 local, the system shall begin a full snapshot of the "
        "primary data volume.",
    )

    _write_cell(t.rows[1].cells[0], "System")
    _write_cell(
        t.rows[1].cells[1],
        "The system must verify snapshot integrity before promoting "
        "the image to the off-site store.",
    )

    # Unattributed row — first cell blank.
    _write_cell(t.rows[2].cells[0], "")
    _write_cell(
        t.rows[2].cells[1],
        "A summary report shall be emailed to the on-call distribution "
        "list once the run completes.",
    )

    _write_cell(t.rows[3].cells[0], "System")
    _write_cell(
        t.rows[3].cells[1],
        "If verification fails, the system must retain the primary "
        "volume in read-write mode and page the on-call engineer.",
    )

    # Another unattributed row, this time with generic content.
    _write_cell(t.rows[4].cells[0], "")
    _write_cell(
        t.rows[4].cells[1],
        "Retention policy shall prune snapshots older than 30 days.",
    )

    doc.save(HERE / "implicit_system_actor.docx")


# ---------------------------------------------------------------------------
# Sample 4 — Heavy passive voice
# ---------------------------------------------------------------------------


def build_passive_voice() -> None:
    """Passive-voice prose: 'shall be performed by X'.

    Primary-column is present but the semantic actor is buried in a
    trailing by-phrase. Regex / heuristic paths should miss the
    by-phrase agent; NLP (dependency parse / NER) should catch most.
    """
    doc = Document()
    doc.add_heading("Procedure \u2014 Change-Management Review", level=0)
    doc.add_paragraph("Document ID: PROC-PASSIVE-004")
    doc.add_paragraph(
        "Review steps for a production change request. Language is "
        "deliberately passive throughout."
    )

    _add_heading(doc, "4. Review Steps", level=1)
    t = doc.add_table(rows=4, cols=2)
    t.style = "Table Grid"

    _write_cell(t.rows[0].cells[0], "Change Board")
    _write_cell(
        t.rows[0].cells[1],
        "Each change request shall be reviewed by the Change Board "
        "before any production deployment.",
    )

    _write_cell(t.rows[1].cells[0], "Change Board")
    _write_cell(
        t.rows[1].cells[1],
        "Approval must be granted by the Security Officer when the "
        "change touches authentication flows.",
    )

    _write_cell(t.rows[2].cells[0], "Change Board")
    _write_cell(
        t.rows[2].cells[1],
        "A rollback plan shall be signed off by the Release Manager "
        "prior to the change window.",
    )

    _write_cell(t.rows[3].cells[0], "Change Board")
    _write_cell(
        t.rows[3].cells[1],
        "Post-change validation must be completed by the Duty "
        "Engineer within 30 minutes of deployment.",
    )

    doc.save(HERE / "passive_voice.docx")


# ---------------------------------------------------------------------------
# Sample 5 — Parallel / interleaved flows
# ---------------------------------------------------------------------------


def build_parallel_flows() -> None:
    """Two actors progress in parallel; their steps interleave.

    Tests row-ordering and actor continuity across a longer table.
    No nesting; the complexity is in the interleaving and the fact
    that each actor's narrative has to be reconstructed from non-
    contiguous rows.
    """
    doc = Document()
    doc.add_heading("Procedure \u2014 Dual-Approver Release", level=0)
    doc.add_paragraph("Document ID: PROC-PARALLEL-005")
    doc.add_paragraph(
        "Two approvers work in parallel during a release window. The "
        "QA Lead exercises the release candidate while the Release "
        "Manager coordinates the deployment pipeline."
    )

    _add_heading(doc, "5. Release Window", level=1)
    t = doc.add_table(rows=8, cols=2)
    t.style = "Table Grid"

    _write_cell(t.rows[0].cells[0], "QA Lead")
    _write_cell(
        t.rows[0].cells[1],
        "The QA Lead shall begin the smoke-test suite against the "
        "release candidate build.",
    )

    _write_cell(t.rows[1].cells[0], "Release Manager")
    _write_cell(
        t.rows[1].cells[1],
        "The Release Manager must stage the deployment package in the "
        "canary environment.",
    )

    _write_cell(t.rows[2].cells[0], "QA Lead")
    _write_cell(
        t.rows[2].cells[1],
        "The QA Lead shall record smoke-test results in the release "
        "tracker within 30 minutes.",
    )

    _write_cell(t.rows[3].cells[0], "Release Manager")
    _write_cell(
        t.rows[3].cells[1],
        "The Release Manager must confirm canary health telemetry for "
        "at least 10 minutes before promotion.",
    )

    _write_cell(t.rows[4].cells[0], "QA Lead")
    _write_cell(
        t.rows[4].cells[1],
        "The QA Lead shall sign off on the release once all smoke "
        "tests pass.",
    )

    _write_cell(t.rows[5].cells[0], "Release Manager")
    _write_cell(
        t.rows[5].cells[1],
        "The Release Manager must promote the canary to production "
        "only after QA sign-off.",
    )

    _write_cell(t.rows[6].cells[0], "QA Lead")
    _write_cell(
        t.rows[6].cells[1],
        "If any smoke test fails, the QA Lead shall halt the release "
        "and notify the Release Manager.",
    )

    _write_cell(t.rows[7].cells[0], "Release Manager")
    _write_cell(
        t.rows[7].cells[1],
        "On halt, the Release Manager must roll back the canary and "
        "capture diagnostic logs for the post-mortem.",
    )

    doc.save(HERE / "parallel_flows.docx")


# ---------------------------------------------------------------------------
# 3-column procedural fixtures (from Eric's 2026-04-23 work-network pass)
#
# Shape:
#     | Actor | Step | Required action |
#
# with a header row whose third column literally reads "Required action".
# Each of the four builders below exercises one of the specific failure
# modes that surfaced on the work-network documents.  They ship with
# paired ``.reqx.yaml`` configs that tell the parser which columns are
# the actor / content (the default 2-column layout wouldn't otherwise
# map onto these).
# ---------------------------------------------------------------------------


def _add_3col_header(t) -> None:
    """Populate the header row of a 3-column procedural table.

    The header is intentionally:

        |  (blank)  |  Step  |  Required Action  |

    — a blank column-1 header, a "Step" column-2 header, and a
    "Required Action" column-3 header (both words capitalised).  The
    combination is the table-type signal Eric flagged in his
    2026-04-23 pass: whenever a 3-column table has this specific
    header shape, every non-empty row should be treated as a
    requirement regardless of modal-keyword content (shall/must/etc).

    All four of the `procedural_*` fixtures below share this header by
    design — they differ in what they stuff into the body rows, not in
    the type signal itself.  The header-aware parser work that
    eventually lands should key off exactly this shape; any other
    3-column table (e.g. "Actor | Step | Required Action") keeps the
    normal keyword-driven detection path.
    """
    _write_cell(t.rows[0].cells[0], "")
    _write_cell(t.rows[0].cells[1], "Step")
    _write_cell(t.rows[0].cells[2], "Required Action")


#: Default per-doc config for the 3-col fixtures below.  Writing the
#: YAML from Python keeps the fixture and its config versioned together
#: and guarantees neither drifts when the generator is re-run.
_3COL_REQX_YAML = (
    "version: 1\n"
    "tables:\n"
    "  actor_column: 1\n"
    "  content_column: 3\n"
    "  min_columns: 3\n"
    "  max_columns: 3\n"
)


def _write_3col_config(stem: str) -> None:
    (HERE / f"{stem}.reqx.yaml").write_text(_3COL_REQX_YAML, encoding="utf-8")


# ---------------------------------------------------------------------------
# Sample 6 — Actor continuation (blank col 1 = inherit from row above)
# ---------------------------------------------------------------------------


def build_procedural_actor_continuation() -> None:
    """3-column procedural table with blank-actor continuation rows.

    Every row is a requirement; column 1 carries the actor except where
    a step continues the previous actor's work and is left blank.  The
    expected behaviour (not yet implemented) is that a blank actor
    cell inherits the nearest non-blank predecessor's actor, so rows
    2 and 4 below should both attribute to the Operator even though
    their own Actor cell is empty.

    Failure mode this targets (FIELD_NOTES §1 + Eric 2026-04-23):
    the current parser treats the blank cell as "unknown actor" and
    either drops the row or attributes it to an empty string, which
    makes the downstream actor summary misleading.
    """
    doc = Document()
    doc.add_heading("Procedure \u2014 Console Shift Handover", level=0)
    doc.add_paragraph("Document ID: PROC-CONT-006")
    doc.add_paragraph(
        "A short procedural table where consecutive steps performed by "
        "the same actor leave the Actor column blank.  The blank cell "
        "continues the actor from the row immediately above."
    )

    _add_heading(doc, "6. Shift Handover Sequence", level=1)
    t = doc.add_table(rows=6, cols=3)
    t.style = "Table Grid"
    _add_3col_header(t)

    _write_cell(t.rows[1].cells[0], "Operator")
    _write_cell(t.rows[1].cells[1], "1")
    _write_cell(
        t.rows[1].cells[2],
        "The Operator shall log in to the control console and confirm "
        "the previous shift's handover notes.",
    )

    # Blank actor cell — continuation from "Operator" above.
    _write_cell(t.rows[2].cells[0], "")
    _write_cell(t.rows[2].cells[1], "2")
    _write_cell(
        t.rows[2].cells[2],
        "Verify that all subsystem status indicators read nominal and "
        "acknowledge any outstanding alarms.",
    )

    _write_cell(t.rows[3].cells[0], "Supervisor")
    _write_cell(t.rows[3].cells[1], "3")
    _write_cell(
        t.rows[3].cells[2],
        "The Supervisor must review the overnight alarm log before "
        "releasing the Operator to normal duties.",
    )

    # Another continuation — this time continuing the Supervisor.
    _write_cell(t.rows[4].cells[0], "")
    _write_cell(t.rows[4].cells[1], "4")
    _write_cell(
        t.rows[4].cells[2],
        "Countersign the handover log once the checks are complete.",
    )

    _write_cell(t.rows[5].cells[0], "Operator")
    _write_cell(t.rows[5].cells[1], "5")
    _write_cell(
        t.rows[5].cells[2],
        "The Operator shall acknowledge receipt of the countersigned "
        "log and begin normal shift duties.",
    )

    doc.save(HERE / "procedural_actor_continuation.docx")
    _write_3col_config("procedural_actor_continuation")


# ---------------------------------------------------------------------------
# Sample 7 — Multi-actor column cell, text picks the real actor
# ---------------------------------------------------------------------------


def build_procedural_multi_actor_cell() -> None:
    """3-column procedural table where column 1 lists multiple candidates.

    Every row names several actors in the Actor column (comma- or
    slash-separated) to indicate "any of these may perform this step".
    The requirement text itself then picks which one actually does the
    work via an explicit subject.  The expected behaviour (not yet
    implemented) is that the parser treats column 1 as a candidate
    *set*, then resolves the concrete primary actor from the sentence
    subject — falling back to the whole set if the text doesn't name
    one explicitly.

    Failure mode this targets: the current parser either concatenates
    the whole cell as one synthetic actor name ("Auth Service, Gateway,
    Logger") or picks the first comma-separated token regardless of
    what the sentence says.
    """
    doc = Document()
    doc.add_heading("Procedure \u2014 Authentication Handshake", level=0)
    doc.add_paragraph("Document ID: PROC-MULTI-007")
    doc.add_paragraph(
        "A handshake with several services participating per step.  "
        "The Actor column enumerates the eligible services; the "
        "Required action column's sentence subject names the one that "
        "actually carries out the step for a given row."
    )

    _add_heading(doc, "7. Handshake Sequence", level=1)
    t = doc.add_table(rows=5, cols=3)
    t.style = "Table Grid"
    _add_3col_header(t)

    _write_cell(t.rows[1].cells[0], "Auth Service, Gateway, Logger")
    _write_cell(t.rows[1].cells[1], "1")
    _write_cell(
        t.rows[1].cells[2],
        "The Gateway shall forward the incoming authentication request "
        "to the Auth Service without modification.",
    )

    _write_cell(t.rows[2].cells[0], "Auth Service, Gateway, Logger")
    _write_cell(t.rows[2].cells[1], "2")
    _write_cell(
        t.rows[2].cells[2],
        "The Auth Service must verify the signed token against the "
        "published public key before any session state is created.",
    )

    # Slash-separated variant — same semantic, different punctuation.
    _write_cell(t.rows[3].cells[0], "Auth Service / Gateway / Logger")
    _write_cell(t.rows[3].cells[1], "3")
    _write_cell(
        t.rows[3].cells[2],
        "The Logger shall emit a structured audit entry containing the "
        "request ID, the verifying service, and the outcome.",
    )

    # Row where the text doesn't explicitly name one of the candidates
    # — resolver should keep the full candidate set as the primary.
    _write_cell(t.rows[4].cells[0], "Auth Service, Gateway, Logger")
    _write_cell(t.rows[4].cells[1], "4")
    _write_cell(
        t.rows[4].cells[2],
        "If the verification fails, the error shall be returned to the "
        "caller with HTTP status 401 and no session token.",
    )

    doc.save(HERE / "procedural_multi_actor_cell.docx")
    _write_3col_config("procedural_multi_actor_cell")


# ---------------------------------------------------------------------------
# Sample 8 — Bullet / numbered list in content cell = multiple requirements
# ---------------------------------------------------------------------------


def build_procedural_bullet_rows() -> None:
    """3-column procedural table where a content cell holds a bullet list.

    Rows whose content is a bulleted or numbered list should emit one
    requirement per bullet — losing that boundary collapses several
    requirements into one synthetic blob and destroys traceability.
    This fixture mixes plain single-sentence rows with list rows so
    tests can pin both paths.

    Failure mode this targets: ``parser._cell_text`` historically
    flattened nested structure within a cell (REVIEW §1.8 — marked
    LOW, partially addressed).  For bullets specifically the boundary
    is still frequently lost; Eric's 2026-04-23 pass flagged this
    explicitly.
    """
    doc = Document()
    doc.add_heading("Procedure \u2014 Pre-Release Checklist", level=0)
    doc.add_paragraph("Document ID: PROC-BULLETS-008")
    doc.add_paragraph(
        "A checklist-style procedure.  Some steps are a single "
        "sentence; others enumerate several sub-actions as a bullet "
        "or numbered list within a single cell.  Each bullet is an "
        "independent requirement."
    )

    _add_heading(doc, "8. Pre-Release Gate", level=1)
    t = doc.add_table(rows=5, cols=3)
    t.style = "Table Grid"
    _add_3col_header(t)

    _write_cell(t.rows[1].cells[0], "Release Manager")
    _write_cell(t.rows[1].cells[1], "1")
    _write_cell(
        t.rows[1].cells[2],
        "The Release Manager shall confirm that the build ID in the "
        "tracker matches the artifact SHA-256 in the pipeline.",
    )

    # Bulleted list row — 3 distinct requirements.
    _write_cell(t.rows[2].cells[0], "QA Lead")
    _write_cell(t.rows[2].cells[1], "2")
    _write_cell(
        t.rows[2].cells[2],
        ("The QA Lead shall execute the following smoke checks:", None),
        ("Run the authentication smoke-test suite and record pass/"
         "fail per case.", "List Bullet"),
        ("Exercise the canary endpoint with a synthetic payload "
         "under nominal load.", "List Bullet"),
        ("Confirm that the error budget dashboard is still green "
         "before clearing the gate.", "List Bullet"),
    )

    _write_cell(t.rows[3].cells[0], "")
    _write_cell(t.rows[3].cells[1], "3")
    _write_cell(
        t.rows[3].cells[2],
        "If any smoke check fails, halt the release and notify the "
        "on-call Release Manager immediately.",
    )

    # Numbered list row — 3 more distinct requirements, different
    # style from the bullet list above.
    _write_cell(t.rows[4].cells[0], "Change Board")
    _write_cell(t.rows[4].cells[1], "4")
    _write_cell(
        t.rows[4].cells[2],
        ("Before granting the final approval the Change Board must:", None),
        ("Review the rollback plan and confirm it executes cleanly "
         "against the staging environment.", "List Number"),
        ("Verify that the Security Officer has signed off when the "
         "change touches authentication flows.", "List Number"),
        ("Record the approval decision and the quorum in the change "
         "register before the change window opens.", "List Number"),
    )

    doc.save(HERE / "procedural_bullet_rows.docx")
    _write_3col_config("procedural_bullet_rows")


# ---------------------------------------------------------------------------
# Sample 9 — Procedural requirements without shall/must keywords
# ---------------------------------------------------------------------------


def build_procedural_no_keywords() -> None:
    """3-column procedural table whose content has no modal keywords.

    Every row is intended to be a requirement — the ``Required action``
    column header is the signal — but the sentences use descriptive
    indicative voice rather than shall/must/should/may.  The current
    keyword-based detector captures none of these rows; the expected
    behaviour (not yet implemented) is a "table-is-requirement-table"
    signal derived from the column header so every non-empty content
    cell becomes a requirement regardless of modal content.

    Failure mode this targets: procedural tables where the author's
    house style uses indicative voice ("The Operator confirms…")
    instead of obligatory voice ("The Operator shall confirm…").
    Without a header-aware signal these documents silently produce
    zero requirements — the worst failure mode because the tool
    appears to have run cleanly.
    """
    doc = Document()
    doc.add_heading("Procedure \u2014 Console Start-of-Shift", level=0)
    doc.add_paragraph("Document ID: PROC-NOKW-009")
    doc.add_paragraph(
        "Written in indicative voice throughout \u2014 no shall / must "
        "/ should / may.  Each row is a requirement by virtue of the "
        "Required action column header, not by virtue of any modal "
        "keyword in the sentence."
    )

    _add_heading(doc, "9. Start-of-Shift Actions", level=1)
    t = doc.add_table(rows=6, cols=3)
    t.style = "Table Grid"
    _add_3col_header(t)

    _write_cell(t.rows[1].cells[0], "Operator")
    _write_cell(t.rows[1].cells[1], "1")
    _write_cell(
        t.rows[1].cells[2],
        "The Operator confirms the console lock status at the start "
        "of every shift.",
    )

    _write_cell(t.rows[2].cells[0], "Operator")
    _write_cell(t.rows[2].cells[1], "2")
    _write_cell(
        t.rows[2].cells[2],
        "The Operator reviews the previous shift's handover notes "
        "before touching any control surface.",
    )

    _write_cell(t.rows[3].cells[0], "Supervisor")
    _write_cell(t.rows[3].cells[1], "3")
    _write_cell(
        t.rows[3].cells[2],
        "The Supervisor validates the active alarm list and "
        "acknowledges any outstanding items.",
    )

    _write_cell(t.rows[4].cells[0], "Supervisor")
    _write_cell(t.rows[4].cells[1], "4")
    _write_cell(
        t.rows[4].cells[2],
        "The Supervisor records the go/no-go decision for the shift "
        "in the shift log.",
    )

    _write_cell(t.rows[5].cells[0], "Operator")
    _write_cell(t.rows[5].cells[1], "5")
    _write_cell(
        t.rows[5].cells[2],
        "The Operator acknowledges the go decision and transitions "
        "the console to normal operating mode.",
    )

    doc.save(HERE / "procedural_no_keywords.docx")
    _write_3col_config("procedural_no_keywords")


# ---------------------------------------------------------------------------
# Sample 10 — Mixed-language procedure
# ---------------------------------------------------------------------------


def build_mixed_language() -> None:
    """2-column table whose content mixes English and Spanish sentences.

    The extractor is keyword-driven and English-focused: modal verbs
    like "shall" / "must" don't have clean Spanish cognates, so
    Spanish-only rows should drop out of the Hard/Soft classifier.
    The English rows must still be captured cleanly.  Useful for
    testing the failure mode where only a subset of the corpus is
    parseable.
    """
    doc = Document()
    doc.add_heading("Procedure \u2014 Customer Onboarding (Bilingual)", level=0)
    doc.add_paragraph("Document ID: PROC-BILING-010")
    doc.add_paragraph(
        "Sample document combining English requirements with Spanish "
        "operator notes.  Only the English rows should be captured as "
        "requirements."
    )

    _add_heading(doc, "10. Onboarding Steps", level=1)
    t = doc.add_table(rows=6, cols=2)
    t.style = "Table Grid"

    _write_cell(t.rows[0].cells[0], "Account Manager")
    _write_cell(
        t.rows[0].cells[1],
        "The Account Manager shall validate the customer identification "
        "before creating any account.",
    )

    _write_cell(t.rows[1].cells[0], "Gestor de Cuentas")
    _write_cell(
        t.rows[1].cells[1],
        "El Gestor de Cuentas debe verificar los documentos del cliente "
        "antes de crear la cuenta.",
    )

    _write_cell(t.rows[2].cells[0], "Compliance Officer")
    _write_cell(
        t.rows[2].cells[1],
        "The Compliance Officer must review the validation result within "
        "two business days.",
    )

    _write_cell(t.rows[3].cells[0], "Oficial de Cumplimiento")
    _write_cell(
        t.rows[3].cells[1],
        "El Oficial de Cumplimiento revisa el resultado de la "
        "validaci\u00f3n en un plazo de dos d\u00edas h\u00e1biles.",
    )

    _write_cell(t.rows[4].cells[0], "Account Manager")
    _write_cell(
        t.rows[4].cells[1],
        "The Account Manager shall notify the customer of the approval "
        "decision via the registered email address.",
    )

    _write_cell(t.rows[5].cells[0], "Gestor de Cuentas")
    _write_cell(
        t.rows[5].cells[1],
        "El Gestor de Cuentas notifica al cliente sobre la decisi\u00f3n "
        "de aprobaci\u00f3n por correo electr\u00f3nico.",
    )

    doc.save(HERE / "mixed_language.docx")


# ---------------------------------------------------------------------------
# Sample 11 — Long procedure (throughput / cancel-path stress)
# ---------------------------------------------------------------------------


def build_long_procedure() -> None:
    """A deliberately long 2-column procedure (50+ rows).

    Exercises the throughput path, the file-progress callback cadence
    in the GUI, and — when paired with a short ``cancel_check`` loop
    in a test — the cancel path's ability to break out mid-table
    without writing a partial output.  Content is industry-generic
    release-engineering prose; actors rotate through a fixed roster
    so the actor-ID path has many samples to lock in.
    """
    doc = Document()
    doc.add_heading("Procedure \u2014 Extended Release Workflow", level=0)
    doc.add_paragraph("Document ID: PROC-LONG-011")
    doc.add_paragraph(
        "Fifty-step release workflow.  Used as a throughput and "
        "cancel-path stress fixture; not a template for a real release "
        "procedure."
    )

    _add_heading(doc, "11. Extended Release Steps", level=1)

    actors = ("QA Lead", "Release Manager", "Duty Engineer", "Security Officer")
    verbs = (
        "shall verify",
        "must confirm",
        "shall record",
        "must notify",
        "shall validate",
        "must sign off",
        "shall capture",
        "must approve",
    )
    subjects = (
        "the build manifest",
        "the release candidate artifact",
        "the canary telemetry dashboard",
        "the rollback plan",
        "the incident response runbook",
        "the change-management ticket",
        "the signed approval record",
        "the post-deployment smoke suite",
        "the security scan report",
        "the compliance attestation",
    )

    n = 52   # 50+ rows; a few more so cancel-path timing tests have margin
    t = doc.add_table(rows=n, cols=2)
    t.style = "Table Grid"
    for i in range(n):
        actor = actors[i % len(actors)]
        verb = verbs[i % len(verbs)]
        subject = subjects[i % len(subjects)]
        _write_cell(t.rows[i].cells[0], actor)
        _write_cell(
            t.rows[i].cells[1],
            f"Step {i + 1}: The {actor} {verb} {subject} "
            f"before the next gate.",
        )

    doc.save(HERE / "long_procedure.docx")


# ---------------------------------------------------------------------------
# Runner
# ---------------------------------------------------------------------------


def main() -> None:
    build_simple_two_actors()
    build_ambiguous_roles()
    build_implicit_system_actor()
    build_passive_voice()
    build_parallel_flows()
    build_procedural_actor_continuation()
    build_procedural_multi_actor_cell()
    build_procedural_bullet_rows()
    build_procedural_no_keywords()
    build_mixed_language()
    build_long_procedure()
    docs = sorted(HERE.glob("*.docx"))
    configs = sorted(HERE.glob("*.reqx.yaml"))
    print(f"Generated {len(docs)} procedure fixtures in {HERE}:")
    for p in docs:
        print(f"  - {p.name}")
    if configs:
        print(f"Paired configs ({len(configs)}):")
        for p in configs:
            print(f"  - {p.name}")


if __name__ == "__main__":
    main()
