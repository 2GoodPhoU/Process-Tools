"""Generate a suite of edge-case .docx files for the extractor's tests.

Produces five samples in this folder:

  nested_tables.docx        — 3-level nested tables, requirements at each depth.
  alphanumeric_sections.docx — 'SR-1.2 / REQ-042 / A.1' style section rows.
  boilerplate_heavy.docx    — real Revision History / Glossary / References /
                              Table of Contents sections followed by requirements.
  wide_table.docx           — 4-column table (ID | Actor | Requirement | Notes).
                              Paired with wide_table.reqx.yaml so the parser
                              picks the right columns.
  noise_prose.docx          — future-tense 'will', 'Note:' / 'Example:' prose,
                              'TBD' markers, negations, empty cells.  Paired
                              with noise_prose.reqx.yaml to exercise content
                              filters and keyword tuning.

Run from the project root:
  python samples/edge_cases/generate.py
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.table import _Cell


HERE = Path(__file__).resolve().parent


# ---------------------------------------------------------------------------
# Small helpers — python-docx has a few rough edges we smooth out here.
# ---------------------------------------------------------------------------


def _write_cell(cell: _Cell, *paragraphs, style: str | None = None) -> None:
    """Replace a cell's content with the given paragraphs.

    ``paragraphs`` is a list of ``(text, style)`` pairs or plain strings.  A
    blank cell is created when ``paragraphs`` is empty.
    """
    # python-docx starts each cell with a single empty paragraph we overwrite.
    first = cell.paragraphs[0]
    first.text = ""
    if not paragraphs:
        return
    for i, p in enumerate(paragraphs):
        text, pstyle = (p, style) if isinstance(p, str) else (p[0], p[1])
        if i == 0:
            target = first
        else:
            target = cell.add_paragraph()
        target.text = text
        if pstyle:
            target.style = pstyle


def _add_heading(doc: Document, text: str, level: int) -> None:
    doc.add_heading(text, level=level)


def _add_para(doc: Document, text: str, style: str | None = None) -> None:
    p = doc.add_paragraph(text)
    if style:
        p.style = style


# ---------------------------------------------------------------------------
# Sample 1 — Deep nested tables (3 levels)
# ---------------------------------------------------------------------------


def build_nested_tables() -> None:
    """Requirement lives at depth 3 inside nested tables.

    Structure:
        Heading 1: "4. Integration Requirements"
        Top-level table (2 cols):
          R1:  [section row] "4.1 Interfaces" | <nested Table A>
            Nested Table A (2 cols):
              R1:  "Ingress Gateway" | <nested Table B + normal req>
                Nested Table B (2 cols):
                  R1: "Ingress (deep)" | "The Ingress Gateway shall drop
                                         malformed packets within 10ms."
              R2:  "Egress Gateway"  | "Egress shall log every outbound
                                       packet with an ISO-8601 timestamp."
          R2: "Flight Computer" | bullet list + paragraph requirements
    """
    doc = Document()

    doc.add_heading("Example System — Integration Specification", level=0)
    doc.add_paragraph("Document ID: SPEC-NESTED-001")
    doc.add_paragraph(
        "This spec demonstrates requirements hidden inside nested tables."
    )

    _add_heading(doc, "4. Integration Requirements", level=1)

    outer = doc.add_table(rows=2, cols=2)
    outer.style = "Table Grid"

    # --- Row 1: section row with two nested tables of requirements ------------
    _write_cell(outer.rows[0].cells[0], "4.1 Interfaces")

    r1c2 = outer.rows[0].cells[1]
    _write_cell(r1c2, "Interfaces between subsystems.")

    nested_a = r1c2.add_table(rows=2, cols=2)
    nested_a.style = "Table Grid"

    # Nested A, Row 1: contains yet another table + a normal paragraph
    _write_cell(nested_a.rows[0].cells[0], "Ingress Gateway")
    a1c2 = nested_a.rows[0].cells[1]
    _write_cell(a1c2, "Gateway handles inbound traffic.")
    nested_b = a1c2.add_table(rows=1, cols=2)
    nested_b.style = "Table Grid"
    _write_cell(nested_b.rows[0].cells[0], "Ingress (deep)")
    _write_cell(
        nested_b.rows[0].cells[1],
        "The Ingress Gateway shall drop malformed packets within 10 ms.",
    )
    # Also add a normal paragraph req in the same cell as the nested table
    a1c2.add_paragraph(
        "The Ingress Gateway must reject traffic from unknown sources."
    )

    # Nested A, Row 2
    _write_cell(nested_a.rows[1].cells[0], "Egress Gateway")
    _write_cell(
        nested_a.rows[1].cells[1],
        "Egress shall log every outbound packet with an ISO-8601 timestamp.",
    )

    # --- Row 2: normal actor row with bullets and prose ----------------------
    _write_cell(outer.rows[1].cells[0], "Flight Computer")
    r2c2 = outer.rows[1].cells[1]
    _write_cell(r2c2, "The Flight Computer shall boot within 5 seconds.")
    r2c2.add_paragraph(
        "Watchdog timeouts must trigger a soft reboot.", style="List Bullet",
    )
    r2c2.add_paragraph(
        "Telemetry may be suspended during reboot.", style="List Bullet",
    )

    doc.save(HERE / "nested_tables.docx")


# ---------------------------------------------------------------------------
# Sample 2 — Alphanumeric section schemes
# ---------------------------------------------------------------------------


def build_alphanumeric_sections() -> None:
    """Section rows use 'SR-1.2', 'REQ-042', 'A.1' etc. to exercise the
    broadened default ``tables.section_prefix`` regex.
    """
    doc = Document()
    doc.add_heading("Alphanumeric Section Sample", level=0)
    doc.add_paragraph("Document ID: SPEC-ALPHA-002")

    _add_heading(doc, "5. Security Requirements", level=1)

    t = doc.add_table(rows=6, cols=2)
    t.style = "Table Grid"

    # Section-style rows interleaved with actor rows, all using non-numeric
    # section prefixes.
    _write_cell(t.rows[0].cells[0], "SR-1.1 Access Control")
    _write_cell(
        t.rows[0].cells[1],
        "Policy rules for role-based access are defined below.",
    )

    _write_cell(t.rows[1].cells[0], "Auth Service")
    _write_cell(
        t.rows[1].cells[1],
        "The Auth Service shall enforce RBAC on every request.",
    )

    _write_cell(t.rows[2].cells[0], "REQ-042 Key Rotation")
    _write_cell(
        t.rows[2].cells[1],
        "Keys are rotated on a fixed cadence.",
    )

    _write_cell(t.rows[3].cells[0], "Key Manager")
    _write_cell(
        t.rows[3].cells[1],
        "The Key Manager shall rotate signing keys every 90 days.",
    )

    _write_cell(t.rows[4].cells[0], "A.1 Annex — Audit")
    _write_cell(
        t.rows[4].cells[1],
        "Audit logs are handled externally to the control plane.",
    )

    _write_cell(t.rows[5].cells[0], "Audit Collector")
    _write_cell(
        t.rows[5].cells[1],
        "The Audit Collector must forward events to the SIEM within 60 "
        "seconds.",
    )

    doc.save(HERE / "alphanumeric_sections.docx")


# ---------------------------------------------------------------------------
# Sample 3 — Boilerplate-heavy doc (to test skip_sections)
# ---------------------------------------------------------------------------


def build_boilerplate_heavy() -> None:
    """Opens with Revision History, TOC, Glossary, and References tables
    (none of which should contribute requirements), then a real section.
    """
    doc = Document()
    doc.add_heading("Boilerplate-Heavy Specification", level=0)
    doc.add_paragraph("Document ID: SPEC-BOILER-003")

    # -- Revision History ------------------------------------------------------
    _add_heading(doc, "1. Revision History", level=1)
    rev = doc.add_table(rows=3, cols=2)
    rev.style = "Table Grid"
    _write_cell(rev.rows[0].cells[0], "Revision History")
    _write_cell(rev.rows[0].cells[1], "Change log")
    _write_cell(rev.rows[1].cells[0], "v1.0")
    _write_cell(
        rev.rows[1].cells[1],
        "Initial release. The review committee shall publish updates "
        "quarterly.",  # <-- this "shall" would be a false positive
    )
    _write_cell(rev.rows[2].cells[0], "v1.1")
    _write_cell(
        rev.rows[2].cells[1],
        "Minor typo fixes. Authors must proofread drafts.",
    )

    # -- Table of Contents -----------------------------------------------------
    _add_heading(doc, "2. Table of Contents", level=1)
    toc = doc.add_table(rows=2, cols=2)
    toc.style = "Table Grid"
    _write_cell(toc.rows[0].cells[0], "Table of Contents")
    _write_cell(toc.rows[0].cells[1], "Page")
    _write_cell(toc.rows[1].cells[0], "Glossary")
    _write_cell(toc.rows[1].cells[1], "See page 3")

    # -- Glossary --------------------------------------------------------------
    _add_heading(doc, "3. Glossary", level=1)
    g = doc.add_table(rows=2, cols=2)
    g.style = "Table Grid"
    _write_cell(g.rows[0].cells[0], "Glossary")
    _write_cell(g.rows[0].cells[1], "Definitions")
    _write_cell(g.rows[1].cells[0], "RBAC")
    _write_cell(
        g.rows[1].cells[1],
        "Role-based access control. Users shall be grouped by role.",
    )

    # -- References ------------------------------------------------------------
    _add_heading(doc, "4. References", level=1)
    r = doc.add_table(rows=2, cols=2)
    r.style = "Table Grid"
    _write_cell(r.rows[0].cells[0], "References")
    _write_cell(r.rows[0].cells[1], "Citations")
    _write_cell(r.rows[1].cells[0], "ISO-27001")
    _write_cell(
        r.rows[1].cells[1],
        "Readers must consult the standard for full details.",
    )

    # -- Real requirements -----------------------------------------------------
    _add_heading(doc, "5. System Requirements", level=1)
    real = doc.add_table(rows=2, cols=2)
    real.style = "Table Grid"
    _write_cell(real.rows[0].cells[0], "Auth Service")
    _write_cell(
        real.rows[0].cells[1],
        "The Auth Service shall log every failed login attempt.",
    )
    _write_cell(real.rows[1].cells[0], "Admin Console")
    _write_cell(
        real.rows[1].cells[1],
        "The Admin Console must support MFA for all administrative users.",
    )

    doc.save(HERE / "boilerplate_heavy.docx")

    # Per-doc config: suppress each boilerplate section's first-column title
    # so none of its content leaks into the requirement output.
    (HERE / "boilerplate_heavy.reqx.yaml").write_text(
        "version: 1\n"
        "skip_sections:\n"
        "  titles:\n"
        "    - Revision History\n"
        "    - Table of Contents\n"
        "    - Glossary\n"
        "    - References\n"
        "    # Each boilerplate table uses its title as the first cell, so\n"
        "    # these matches are enough to suppress the entire table's rows.\n"
        "    - v1.0\n"
        "    - v1.1\n"
        "    - RBAC\n"
        "    - ISO-27001\n",
        encoding="utf-8",
    )


# ---------------------------------------------------------------------------
# Sample 4 — Wide 4-column table (ID | Actor | Requirement | Notes)
# ---------------------------------------------------------------------------


def build_wide_table() -> None:
    """Tests column remapping: actor lives in col 2, content in col 3."""
    doc = Document()
    doc.add_heading("Wide Table Sample", level=0)
    doc.add_paragraph("Document ID: SPEC-WIDE-004")

    _add_heading(doc, "6. Control System", level=1)
    t = doc.add_table(rows=4, cols=4)
    t.style = "Table Grid"

    # Header row
    _write_cell(t.rows[0].cells[0], "ID")
    _write_cell(t.rows[0].cells[1], "Actor")
    _write_cell(t.rows[0].cells[2], "Requirement")
    _write_cell(t.rows[0].cells[3], "Notes")

    _write_cell(t.rows[1].cells[0], "R-001")
    _write_cell(t.rows[1].cells[1], "Controller")
    _write_cell(
        t.rows[1].cells[2],
        "The Controller shall sample sensor data at 100 Hz.",
    )
    _write_cell(t.rows[1].cells[3], "Nominal rate")

    _write_cell(t.rows[2].cells[0], "R-002")
    _write_cell(t.rows[2].cells[1], "Actuator")
    _write_cell(
        t.rows[2].cells[2],
        "The Actuator must respond to commands within 20 ms.",
    )
    _write_cell(t.rows[2].cells[3], "Measured end-to-end")

    _write_cell(t.rows[3].cells[0], "R-003")
    _write_cell(t.rows[3].cells[1], "Operator Console")
    _write_cell(
        t.rows[3].cells[2],
        "The Operator Console should display alarms within 1 s of onset.",
    )
    _write_cell(t.rows[3].cells[3], "Soft real-time")

    doc.save(HERE / "wide_table.docx")

    # Per-doc config so the parser knows this is a 4-col doc with actor in
    # column 2 and content in column 3.  Also drops the built-in 'requirement'
    # noun keyword so the table's own header row ("Requirement") doesn't get
    # classified as a requirement itself.
    (HERE / "wide_table.reqx.yaml").write_text(
        "version: 1\n"
        "tables:\n"
        "  actor_column: 2\n"
        "  content_column: 3\n"
        "  min_columns: 4\n"
        "  max_columns: 4\n"
        "keywords:\n"
        "  hard_remove: [requirement, requirements]\n",
        encoding="utf-8",
    )


# ---------------------------------------------------------------------------
# Sample 5 — Noise prose (content filters + keyword tuning)
# ---------------------------------------------------------------------------


def build_noise_prose() -> None:
    """Exercises content.skip_if_starts_with, skip_pattern, and keyword
    tuning.  Paired with noise_prose.reqx.yaml.
    """
    doc = Document()
    doc.add_heading("Noise-Prose Sample", level=0)
    doc.add_paragraph("Document ID: SPEC-NOISE-005")

    # Preamble with a future-tense "will" that should NOT be a requirement
    # once the config's ``keywords.hard_remove: [will]`` applies.
    doc.add_paragraph(
        "This document will serve as the authoritative reference for the "
        "subsystem."
    )

    _add_heading(doc, "7. System Behaviour", level=1)

    t = doc.add_table(rows=6, cols=2)
    t.style = "Table Grid"

    _write_cell(t.rows[0].cells[0], "Scheduler")
    sch = t.rows[0].cells[1]
    _write_cell(
        sch,
        "The Scheduler shall dispatch tasks every 100 ms.",
        "Note: this behaviour is informational and does not constitute a "
        "requirement.",
        "Example: dispatch happens on the tick boundary.",
        "The Scheduler must not preempt a critical task.",  # negation
    )

    _write_cell(t.rows[1].cells[0], "Logger")
    _write_cell(
        t.rows[1].cells[1],
        "The Logger will attempt to flush buffers every second.",
        # "will" only sentence → dropped when keywords.hard_remove includes
        # "will".
    )

    _write_cell(t.rows[2].cells[0], "Orchestrator")
    _write_cell(
        t.rows[2].cells[1],
        "The Orchestrator shall coordinate all subsystems.",
        "Caution: behaviour during brownout is TBD.",  # TBD matches skip_pattern
    )

    # Empty row (no content, should produce nothing).
    _write_cell(t.rows[3].cells[0], "Monitor")
    _write_cell(t.rows[3].cells[1])  # intentionally empty

    # Whitespace-only cell.
    _write_cell(t.rows[4].cells[0], "Reporter")
    _write_cell(t.rows[4].cells[1], "   ", "\t\t")

    # Short requirement ("Shall do.") — tests confidence='Low' path.
    _write_cell(t.rows[5].cells[0], "Beacon")
    _write_cell(t.rows[5].cells[1], "Shall ping.")

    doc.save(HERE / "noise_prose.docx")

    (HERE / "noise_prose.reqx.yaml").write_text(
        "version: 1\n"
        "keywords:\n"
        "  hard_remove: [will]\n"
        "content:\n"
        "  skip_if_starts_with:\n"
        "    - \"Note:\"\n"
        "    - \"Example:\"\n"
        "    - \"Caution:\"\n"
        "  skip_pattern: '\\bTBD\\b'\n",
        encoding="utf-8",
    )


# ---------------------------------------------------------------------------
# Runner
# ---------------------------------------------------------------------------


def main() -> None:
    build_nested_tables()
    build_alphanumeric_sections()
    build_boilerplate_heavy()
    build_wide_table()
    build_noise_prose()
    print(f"Generated 5 samples in {HERE}:")
    for p in sorted(HERE.glob("*.docx")):
        print(f"  - {p.name}")
    for p in sorted(HERE.glob("*.reqx.yaml")):
        print(f"  - {p.name}")


if __name__ == "__main__":
    main()
