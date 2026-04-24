"""Tests for legacy-format support (REVIEW §3.1).

The actual conversion helpers (``convert_doc_to_docx``,
``convert_pdf_to_docx``) are behaviour-tested only when their
dependencies are installed — both are guarded by skipIf so the test
suite stays green on machines without LibreOffice or pdfplumber.

The parts that run unconditionally:

* Discovery / capability checks (``find_soffice``, ``has_soffice``).
* The ``prepare_for_parser`` context-manager routing for each
  supported extension.
* The error-path messages — we want user-facing guidance to say
  "install LibreOffice" rather than a bare stack trace.

Run:  python -m unittest tests.test_legacy_formats
"""

from __future__ import annotations

import os
import tempfile
import unittest
from pathlib import Path
from unittest import mock

from requirements_extractor.legacy_formats import (
    LibreOfficeUnavailable,
    PdfSupportUnavailable,
    _missing_libreoffice_message,
    _missing_pdfplumber_message,
    convert_doc_to_docx,
    convert_pdf_to_docx,
    find_soffice,
    has_soffice,
    prepare_for_parser,
)


ROOT = Path(__file__).resolve().parent.parent
SAMPLE_DOCX = ROOT / "samples" / "sample_spec.docx"


_SOFFICE_AVAILABLE = has_soffice()
try:  # pragma: no cover — availability-gated
    import pdfplumber  # type: ignore  # noqa: F401
    _PDFPLUMBER_AVAILABLE = True
except ImportError:
    _PDFPLUMBER_AVAILABLE = False


# ---------------------------------------------------------------------------
# Capability discovery
# ---------------------------------------------------------------------------


class TestDiscovery(unittest.TestCase):
    def test_find_soffice_returns_str_or_none(self) -> None:
        result = find_soffice()
        self.assertTrue(result is None or isinstance(result, str))

    def test_has_soffice_is_consistent(self) -> None:
        # has_soffice() is a Bool wrapper over find_soffice().
        self.assertEqual(has_soffice(), find_soffice() is not None)

    def test_find_soffice_returns_none_when_nothing_on_path(self) -> None:
        """With PATH scrubbed and all platform fallbacks missing,
        find_soffice() must return None rather than leak an exception."""
        with mock.patch("shutil.which", return_value=None), \
             mock.patch("os.path.exists", return_value=False):
            self.assertIsNone(find_soffice())


# ---------------------------------------------------------------------------
# Error messages are user-facing — pin their shape so a reviewer or a
# translator-to-other-locales can see what ships.
# ---------------------------------------------------------------------------


class TestErrorMessages(unittest.TestCase):
    def test_libreoffice_message_mentions_install(self) -> None:
        msg = _missing_libreoffice_message()
        self.assertIn("LibreOffice", msg)
        # Platform-specific hint should appear on at least one of the
        # three paths (Windows / macOS / Linux).
        self.assertTrue(
            "libreoffice.org" in msg.lower()
            or "brew" in msg.lower()
            or "apt" in msg.lower()
            or "dnf" in msg.lower()
            or "pacman" in msg.lower()
        )

    def test_pdfplumber_message_mentions_pip(self) -> None:
        msg = _missing_pdfplumber_message()
        self.assertIn("pdfplumber", msg)
        self.assertIn("pip install", msg)


# ---------------------------------------------------------------------------
# prepare_for_parser — routing by extension
# ---------------------------------------------------------------------------


class TestPrepareForParserRouting(unittest.TestCase):
    """The context manager should:
      * Return the input path untouched for ``.docx``.
      * Convert ``.doc`` / ``.pdf`` into a temp ``.docx`` when the
        capability is available; otherwise raise a friendly error.
      * Reject unknown extensions.
    """

    def test_docx_passthrough_no_tempdir(self) -> None:
        """Fast path: .docx inputs should never go through a tempdir."""
        if not SAMPLE_DOCX.exists():
            self.skipTest("sample_spec.docx missing")
        with prepare_for_parser(SAMPLE_DOCX) as ready:
            self.assertEqual(ready, SAMPLE_DOCX)
            # Ready path is the original, not a temp copy.
            self.assertTrue(ready.exists())

    def test_unsupported_extension_raises(self) -> None:
        with tempfile.TemporaryDirectory() as d:
            weird = Path(d) / "foo.rtf"
            weird.write_text("nope")
            with self.assertRaises(ValueError) as ctx:
                with prepare_for_parser(weird):
                    pass
            self.assertIn(".rtf", str(ctx.exception))
            self.assertIn(".docx", str(ctx.exception))

    def test_doc_without_libreoffice_raises_friendly(self) -> None:
        """With soffice discovery returning None, a .doc input must
        surface the install-instruction message rather than a vague
        'file not found' or subprocess error."""
        with tempfile.TemporaryDirectory() as d:
            fake_doc = Path(d) / "legacy.doc"
            # python-docx can't read this; we won't get that far.
            fake_doc.write_bytes(b"\xD0\xCF\x11\xE0\xA1\xB1\x1A\xE1")
            with mock.patch(
                "requirements_extractor.legacy_formats.find_soffice",
                return_value=None,
            ):
                with self.assertRaises(LibreOfficeUnavailable) as ctx:
                    with prepare_for_parser(fake_doc):
                        pass
                self.assertIn("LibreOffice", str(ctx.exception))

    def test_pdf_without_pdfplumber_raises_friendly(self) -> None:
        """With pdfplumber import blocked, a .pdf input must surface
        the install hint rather than a bare ImportError trace."""
        with tempfile.TemporaryDirectory() as d:
            fake_pdf = Path(d) / "spec.pdf"
            fake_pdf.write_bytes(b"%PDF-1.4\n%bogus")
            # Shadow pdfplumber in sys.modules so the import inside
            # convert_pdf_to_docx fails even on machines where it's
            # really installed.
            import sys
            with mock.patch.dict(sys.modules, {"pdfplumber": None}):
                with self.assertRaises(PdfSupportUnavailable) as ctx:
                    with prepare_for_parser(fake_pdf):
                        pass
                self.assertIn("pdfplumber", str(ctx.exception))


# ---------------------------------------------------------------------------
# .doc conversion — runs only if LibreOffice is actually installed
# ---------------------------------------------------------------------------


@unittest.skipUnless(
    _SOFFICE_AVAILABLE,
    "LibreOffice not installed on this machine — skipping .doc conversion test.",
)
class TestDocConversionIntegration(unittest.TestCase):  # pragma: no cover
    """Integration test: build a tiny .doc via LibreOffice itself,
    convert it back to .docx, confirm the result opens cleanly.

    Gated — this test only runs on machines with LibreOffice installed
    so the suite stays green in restricted environments (work network,
    bare-python CI).
    """

    def test_roundtrip_converts_to_docx(self) -> None:
        # Build a source .docx we can round-trip.
        from docx import Document
        with tempfile.TemporaryDirectory() as d:
            src = Path(d) / "src.docx"
            doc = Document()
            doc.add_paragraph("Hello from a .doc round-trip test.")
            doc.save(str(src))
            # Step 1: .docx -> .doc (via soffice)
            import subprocess
            subprocess.run(
                [find_soffice(), "--headless", "--convert-to", "doc",
                 "--outdir", d, str(src)],
                check=True, timeout=120,
                stdout=subprocess.PIPE, stderr=subprocess.PIPE,
            )
            legacy = Path(d) / "src.doc"
            self.assertTrue(legacy.exists(), "soffice did not produce a .doc")
            # Step 2: .doc -> .docx (our function under test)
            out_dir = Path(d) / "out"
            converted = convert_doc_to_docx(legacy, out_dir)
            self.assertTrue(converted.exists())
            self.assertEqual(converted.suffix, ".docx")
            # Re-open to confirm it's a valid docx.
            Document(str(converted))


# ---------------------------------------------------------------------------
# .pdf conversion — runs only when pdfplumber is importable
# ---------------------------------------------------------------------------


@unittest.skipUnless(
    _PDFPLUMBER_AVAILABLE,
    "pdfplumber not installed — skipping .pdf conversion test.",
)
class TestPdfConversionIntegration(unittest.TestCase):  # pragma: no cover
    """Build a tiny PDF in memory, convert to docx via our helper,
    confirm the output opens and contains the expected table content.
    """

    def _build_tiny_pdf(self, path: Path) -> None:
        # Generate a minimal PDF without any extra reportlab / weasyprint
        # dependency by writing a handcrafted one-page PDF with a
        # simple text line.  This exercises the prose-extraction path;
        # the table path benefits from fixtures but is incidentally
        # exercised by the integration test on a real PDF.
        path.write_bytes(
            b"%PDF-1.4\n"
            b"1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj\n"
            b"2 0 obj<</Type/Pages/Count 1/Kids[3 0 R]>>endobj\n"
            b"3 0 obj<</Type/Page/Parent 2 0 R/MediaBox[0 0 200 200]"
            b"/Contents 4 0 R/Resources<</Font<</F1 5 0 R>>>>>>endobj\n"
            b"4 0 obj<</Length 44>>stream\n"
            b"BT /F1 12 Tf 10 100 Td (Hello PDF) Tj ET\n"
            b"endstream endobj\n"
            b"5 0 obj<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>endobj\n"
            b"xref\n0 6\n0000000000 65535 f \n0000000009 00000 n \n"
            b"0000000052 00000 n \n0000000098 00000 n \n"
            b"0000000183 00000 n \n0000000252 00000 n \ntrailer\n"
            b"<</Size 6/Root 1 0 R>>\nstartxref\n313\n%%EOF\n"
        )

    def test_prose_pdf_produces_docx(self) -> None:
        from docx import Document
        with tempfile.TemporaryDirectory() as d:
            pdf = Path(d) / "tiny.pdf"
            self._build_tiny_pdf(pdf)
            out_dir = Path(d) / "out"
            result = convert_pdf_to_docx(pdf, out_dir)
            self.assertTrue(result.exists())
            self.assertEqual(result.suffix, ".docx")
            # Document opens cleanly.
            docx = Document(str(result))
            # Expect at least one paragraph containing our text.
            texts = [p.text for p in docx.paragraphs]
            joined = "\n".join(texts)
            self.assertIn("Hello PDF", joined)


# ---------------------------------------------------------------------------
# PDF no-double-emit regression (from the 2026-04-24 test-session finding)
# ---------------------------------------------------------------------------


@unittest.skipUnless(
    _SOFFICE_AVAILABLE and _PDFPLUMBER_AVAILABLE,
    "Needs both LibreOffice (to build a test PDF) and pdfplumber "
    "(to read it). Skipping when either is absent.",
)
class TestPdfNoDoubleEmit(unittest.TestCase):  # pragma: no cover
    """When a PDF page has tables, the extracted-text pass must be
    suppressed so each requirement appears exactly once in the output
    (once per table row) rather than twice (table row + preamble).

    Integration test: take one of the checked-in .docx fixtures,
    convert it to PDF via soffice, run it through convert_pdf_to_docx,
    then run it through the main parser and count the emitted
    requirements.  The count must match the original .docx's count.
    """

    def test_pdf_of_simple_two_actors_emits_no_duplicates(self) -> None:
        import subprocess
        from requirements_extractor.parser import parse_docx
        from requirements_extractor.config import resolve_config

        src = ROOT / "samples" / "procedures" / "simple_two_actors.docx"
        if not src.exists():
            self.skipTest("simple_two_actors fixture missing")

        # Baseline: parse the original .docx directly.
        baseline = parse_docx(
            src, resolver_fn=lambda t, a: [],
            config=resolve_config(docx_path=src),
        )

        with tempfile.TemporaryDirectory() as d:
            td = Path(d)
            # .docx -> .pdf via soffice.
            subprocess.run(
                [find_soffice(), "--headless", "--convert-to", "pdf",
                 "--outdir", str(td), str(src)],
                check=True, timeout=120,
                stdout=subprocess.PIPE, stderr=subprocess.PIPE,
            )
            pdf = td / (src.stem + ".pdf")
            self.assertTrue(pdf.exists(), "soffice didn't produce a PDF")

            # .pdf -> synthetic .docx via our converter.
            out_dir = td / "out"
            synthetic = convert_pdf_to_docx(pdf, out_dir)

            # Parse the synthetic .docx through the main pipeline and
            # compare the requirement count.  The PDF path is lossier
            # than the .docx baseline in general (structure can drift),
            # but the invariant we pin here is that the count must NOT
            # be 2x the baseline — which is what double-emit looked like
            # before the fix.
            through_pdf = parse_docx(
                synthetic, resolver_fn=lambda t, a: [],
                config=resolve_config(docx_path=synthetic),
            )
            # Must not be double-counted.  Allow equal or lower (PDF
            # conversion may drop some rows if a cell wraps awkwardly,
            # but it must never produce MORE than the baseline).
            self.assertLessEqual(
                len(through_pdf), len(baseline),
                msg=(
                    f"PDF path produced {len(through_pdf)} requirements; "
                    f"baseline .docx produced {len(baseline)}. "
                    f"More rows from the PDF than the .docx means the "
                    f"duplicate-emit regression is back."
                ),
            )
            # And a strictly positive sanity check — the PDF shouldn't
            # end up with zero requirements (that would mean the table
            # extraction broke).
            self.assertGreater(
                len(through_pdf), 0,
                msg="PDF path emitted zero requirements — table extraction broke.",
            )


if __name__ == "__main__":
    unittest.main()
