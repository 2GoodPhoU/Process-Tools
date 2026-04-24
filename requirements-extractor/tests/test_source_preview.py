"""Tests for REVIEW §3.8 — inline source-preview / Context column.

The "Context" column carries the surrounding source text so a reviewer
can sanity-check a requirement without opening the original .docx.  The
implementation is split across three points:

* ``parser._build_context`` — the snippet builder (whitespace collapse,
  redundancy suppression, length cap with sentence-friendly truncation).
* ``parser._emit_candidate`` — accepts the ``context`` kwarg and feeds
  it to the Requirement constructor.
* ``writer.COLUMNS`` — appends a "Context" column to the xlsx output.

These tests pin all three so the column can't silently regress.
"""

from __future__ import annotations

import tempfile
import unittest
from pathlib import Path

from docx import Document
from openpyxl import load_workbook

from requirements_extractor.extractor import extract_from_files
from requirements_extractor.models import Requirement
from requirements_extractor.parser import _MAX_CONTEXT_CHARS, _build_context
from requirements_extractor.writer import COLUMNS


class TestBuildContext(unittest.TestCase):
    """Pure-function tests for the snippet builder."""

    def test_empty_context_returns_empty(self):
        self.assertEqual(_build_context("", "anything"), "")

    def test_redundant_context_collapses_to_empty(self):
        # When the surrounding text IS the requirement text (one-sentence
        # paragraph), there's no value-add in carrying it twice.
        text = "The system shall log every request."
        self.assertEqual(_build_context(text, text), "")

    def test_redundant_after_whitespace_collapse(self):
        # Stray double-spaces and trailing newline shouldn't defeat the
        # redundancy check — both sides normalise the same way.
        self.assertEqual(
            _build_context(
                "  The system   shall log every request.  \n",
                "The system shall log every request.",
            ),
            "",
        )

    def test_redundant_after_casefold(self):
        # ALL CAPS surrounding paragraph against title-case requirement
        # is still redundant for review purposes.
        self.assertEqual(
            _build_context(
                "THE SYSTEM SHALL LOG EVERY REQUEST.",
                "The system shall log every request.",
            ),
            "",
        )

    def test_short_context_passes_through_collapsed(self):
        # Multi-line input collapses to single-line; under-cap so no truncation.
        ctx = "Authentication is critical.\nThe system shall verify.\nFailures get logged."
        text = "The system shall verify."
        result = _build_context(ctx, text)
        # Whitespace-collapsed so xlsx cell stays single-line.
        self.assertEqual(
            result,
            "Authentication is critical. The system shall verify. Failures get logged.",
        )

    def test_long_context_truncates_at_word_boundary_with_ellipsis(self):
        # Build a context that's well over the cap and verify the cut
        # lands at a space, not mid-word.
        ctx = "alpha bravo charlie delta " * 50  # >> 280 chars
        text = "the system shall do something else"
        result = _build_context(ctx, text)
        self.assertLessEqual(len(result), _MAX_CONTEXT_CHARS)
        self.assertTrue(result.endswith("…"), f"missing ellipsis: {result!r}")
        # No mid-word cut: the char immediately before the ellipsis is
        # the end of a word (since we always cut at a space).
        body = result[:-1].rstrip()
        # Last word in body should be a complete word from the input.
        last_word = body.split()[-1]
        self.assertIn(last_word, ctx.split())

    def test_truncation_falls_back_when_no_whitespace_in_window(self):
        # Pathological input: no spaces in the first MAX_CONTEXT_CHARS.
        # We still produce something under the cap with the ellipsis
        # rather than throwing.
        ctx = "x" * (_MAX_CONTEXT_CHARS * 2)
        result = _build_context(ctx, "some other text")
        self.assertLessEqual(len(result), _MAX_CONTEXT_CHARS)
        self.assertTrue(result.endswith("…"))


class TestRequirementModelHasContext(unittest.TestCase):
    """The dataclass field exists with the right default."""

    def test_default_context_is_empty_string(self):
        req = Requirement(
            order=1,
            source_file="x.docx",
            heading_trail="",
            section_topic="",
            row_ref="Table 1, Row 1",
            block_ref="Paragraph 1",
            primary_actor="System",
            secondary_actors=[],
            text="The system shall do X.",
            req_type="Hard",
            keywords=["shall"],
            confidence="High",
        )
        self.assertEqual(req.context, "")

    def test_context_field_round_trips(self):
        req = Requirement(
            order=1,
            source_file="x.docx",
            heading_trail="",
            section_topic="",
            row_ref="Table 1, Row 1",
            block_ref="Paragraph 1",
            primary_actor="System",
            secondary_actors=[],
            text="The system shall do X.",
            req_type="Hard",
            keywords=["shall"],
            confidence="High",
            context="Some surrounding paragraph context.",
        )
        self.assertEqual(req.context, "Some surrounding paragraph context.")


class TestWriterContextColumn(unittest.TestCase):
    """The Context column lives at the right place in the workbook."""

    def test_context_is_the_rightmost_column(self):
        # Adding the column at the end means existing column indexing
        # in user formulas / scripts doesn't shift.
        last_header, _width = COLUMNS[-1]
        self.assertEqual(last_header, "Context")

    def test_no_existing_column_shifted(self):
        # The 15 columns that existed before §3.8 must still be in
        # the same order.  Pinned by name to catch accidental reorder.
        names = [name for name, _w in COLUMNS]
        self.assertEqual(
            names,
            [
                "#", "ID", "Source File", "Heading Trail", "Section / Topic",
                "Row Ref", "Block Ref", "Primary Actor", "Secondary Actors",
                "Requirement", "Type", "Polarity", "Keywords", "Confidence",
                "Notes", "Context",
            ],
        )


def _build_multi_sentence_doc(path: Path) -> None:
    """Helper: emit a .docx with a multi-sentence paragraph in a 2-col table.

    Three of the four sentences contain a HARD modal so the extractor
    will produce three Requirement rows from the same paragraph — every
    one of them should carry the surrounding paragraph as context.
    """
    doc = Document()
    doc.add_heading("Test Spec", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Auth Service"
    table.rows[0].cells[1].text = (
        "Authentication is critical for our system. "
        "The system shall verify all credentials. "
        "Failed attempts must be logged. "
        "Audit logs must be retained for at least 90 days."
    )
    doc.save(str(path))


class TestEndToEndContextPopulated(unittest.TestCase):
    """Integration: multi-sentence paragraph → context column populated."""

    def test_xlsx_context_column_present_and_populated(self):
        with tempfile.TemporaryDirectory() as d:
            src = Path(d) / "ctx.docx"
            out = Path(d) / "out.xlsx"
            _build_multi_sentence_doc(src)
            extract_from_files([src], out)
            wb = load_workbook(str(out))
            ws = wb["Requirements"]
            headers = [c.value for c in ws[1]]
            self.assertEqual(headers[-1], "Context")
            ctx_col_idx = len(headers)  # 1-based
            # Every body row should have non-empty context for this fixture
            # (3 hard sentences in a 4-sentence paragraph).
            data_rows = list(ws.iter_rows(min_row=2, values_only=True))
            self.assertGreaterEqual(len(data_rows), 3)
            for row in data_rows:
                ctx_value = row[ctx_col_idx - 1]
                self.assertTrue(
                    ctx_value and "Authentication is critical" in ctx_value,
                    f"expected paragraph as context, got: {ctx_value!r}",
                )

    def test_single_sentence_paragraph_omits_context(self):
        # When a paragraph has only one sentence, context == requirement
        # (post-normalisation) and `_build_context` suppresses it.
        doc = Document()
        doc.add_heading("Test Spec", level=1)
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        table.rows[0].cells[0].text = "System"
        table.rows[0].cells[1].text = "The system shall log every request."
        with tempfile.TemporaryDirectory() as d:
            src = Path(d) / "single.docx"
            out = Path(d) / "out.xlsx"
            doc.save(str(src))
            extract_from_files([src], out)
            wb = load_workbook(str(out))
            ws = wb["Requirements"]
            data_rows = list(ws.iter_rows(min_row=2, values_only=True))
            self.assertEqual(len(data_rows), 1)
            # Context column is the rightmost; should be None or "".
            self.assertFalse(data_rows[0][-1])


if __name__ == "__main__":
    unittest.main()
