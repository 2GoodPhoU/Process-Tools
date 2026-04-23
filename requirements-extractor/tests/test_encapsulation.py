"""Tests for REVIEW §2.4 + §2.7 — encapsulation fixes.

§2.4 gave ``ActorResolver`` a public API (``has_nlp``, ``iter_regex_hits``,
``iter_nlp_hits``, ``iter_matches``) so callers no longer have to reach
into ``_nlp`` / ``_actor_re`` / ``_alias_to_canonical``.

§2.7 wrapped the remaining python-docx private-attribute accesses
(``_Cell._tc``, ``Paragraph._p``) in ``_cell_element`` / ``_paragraph_element``
helpers that will keep working if python-docx adds a public attribute.

These tests pin the behaviour so future refactors don't silently regress
the contract.
"""

from __future__ import annotations

import unittest

from requirements_extractor.actors import ActorEntry, ActorResolver


class ActorResolverPublicApiTests(unittest.TestCase):
    def test_has_nlp_false_by_default(self):
        resolver = ActorResolver([])
        self.assertFalse(resolver.has_nlp())

    def test_iter_regex_hits_yields_canonical_names(self):
        resolver = ActorResolver([
            ActorEntry(name="Cashier", aliases=["teller"]),
            ActorEntry(name="Customer", aliases=[]),
        ])
        text = "The teller greets the customer."
        hits = list(resolver.iter_regex_hits(text))
        self.assertIn("Cashier", hits)
        self.assertIn("Customer", hits)

    def test_iter_regex_hits_excludes_primary(self):
        resolver = ActorResolver([
            ActorEntry(name="Cashier", aliases=[]),
            ActorEntry(name="Customer", aliases=[]),
        ])
        text = "The cashier greets the customer."
        hits = list(resolver.iter_regex_hits(text, primary="Cashier"))
        self.assertEqual(hits, ["Customer"])

    def test_iter_regex_hits_deduplicates(self):
        resolver = ActorResolver([ActorEntry(name="Cashier", aliases=["teller"])])
        text = "The cashier tells the teller that the cashier is busy."
        hits = list(resolver.iter_regex_hits(text))
        self.assertEqual(hits, ["Cashier"])

    def test_iter_regex_hits_empty_when_no_actors(self):
        resolver = ActorResolver([])
        self.assertEqual(list(resolver.iter_regex_hits("anything")), [])

    def test_iter_regex_hits_empty_text(self):
        resolver = ActorResolver([ActorEntry(name="Cashier", aliases=[])])
        self.assertEqual(list(resolver.iter_regex_hits("")), [])

    def test_iter_matches_attributes_source(self):
        resolver = ActorResolver([ActorEntry(name="Cashier", aliases=[])])
        text = "The cashier approves the payment."
        matches = list(resolver.iter_matches(text))
        self.assertEqual(matches, [("Cashier", "regex")])

    def test_iter_nlp_hits_empty_without_nlp(self):
        resolver = ActorResolver([])
        self.assertEqual(list(resolver.iter_nlp_hits("Alice met Bob.")), [])

    def test_resolve_uses_public_api(self):
        """``resolve`` should still work and be consistent with iter_matches."""
        resolver = ActorResolver([
            ActorEntry(name="Cashier", aliases=[]),
            ActorEntry(name="Customer", aliases=[]),
        ])
        text = "The cashier greets the customer and refunds the customer."
        self.assertEqual(
            resolver.resolve(text, primary="Cashier"),
            ["Customer"],
        )


class ParserPrivateAttrWrappersTests(unittest.TestCase):
    """Smoke-test that the python-docx wrapper helpers work on real objects."""

    def test_cell_and_paragraph_element_helpers(self):
        from docx import Document  # python-docx

        from requirements_extractor.parser import (
            _cell_element,
            _paragraph_element,
        )

        doc = Document()
        doc.add_paragraph("Hello")
        table = doc.add_table(rows=1, cols=1)
        cell = table.rows[0].cells[0]
        cell.text = "World"
        paragraph = doc.paragraphs[0]

        # Each helper must return a non-None lxml element whose tag is the
        # expected docx element name.
        tc = _cell_element(cell)
        self.assertIsNotNone(tc)
        self.assertTrue(tc.tag.endswith("}tc"))

        p = _paragraph_element(paragraph)
        self.assertIsNotNone(p)
        self.assertTrue(p.tag.endswith("}p"))

    def test_paragraph_element_raises_clear_error_on_foreign_object(self):
        from requirements_extractor.parser import _paragraph_element

        class NotAParagraph:
            pass

        with self.assertRaises(AttributeError):
            _paragraph_element(NotAParagraph())


if __name__ == "__main__":
    unittest.main()
