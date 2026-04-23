"""Tests for the module-size / de-duplication refactors.

* ``cli.SUBCOMMAND_NAMES`` is the single source of truth for the names
  accepted by the argparse subparser.  ``extract.py`` re-exports it so
  the backward-compat shim automatically tracks new subcommands.  This
  test pins that they stay in sync with what ``build_parser`` actually
  registers.

* ``cli._harvest_auto_actors`` was extracted from ``_run_requirements``
  so the auto-actors pre-pass is testable in isolation and the parent
  function stays legible.  Smoke-test it against a minimal docx.
"""

from __future__ import annotations

import tempfile
import unittest
from pathlib import Path
from types import SimpleNamespace

from docx import Document

from requirements_extractor import cli
from requirements_extractor.cli import SUBCOMMAND_NAMES, build_parser


class SubcommandCatalogueTests(unittest.TestCase):
    def test_catalogue_matches_argparse_registration(self):
        """Every name in SUBCOMMAND_NAMES must be a real subcommand or alias."""
        parser = build_parser()
        # Pull the subparsers action out so we can introspect it.
        subparsers = next(
            a for a in parser._subparsers._group_actions  # noqa: SLF001
            if isinstance(a, type(parser._subparsers._group_actions[0]))  # noqa: SLF001
        )
        registered = set(subparsers.choices.keys())
        # Every catalogue name should be something argparse knows about.
        self.assertTrue(
            SUBCOMMAND_NAMES.issubset(registered),
            f"SUBCOMMAND_NAMES has entries not in argparse: "
            f"{SUBCOMMAND_NAMES - registered}",
        )
        # And every canonical/alias argparse registered should be in the
        # catalogue — so adding a new subcommand forces the author to
        # update the constant.
        self.assertTrue(
            registered.issubset(SUBCOMMAND_NAMES),
            f"argparse knows subcommands the catalogue doesn't: "
            f"{registered - SUBCOMMAND_NAMES}",
        )

    def test_shim_imports_the_constant(self):
        """extract.py must re-use SUBCOMMAND_NAMES, not hand-maintain its own."""
        import extract as shim

        self.assertIs(shim._KNOWN_SUBCOMMANDS, SUBCOMMAND_NAMES)


def _write_minimal_docx(path: Path) -> None:
    d = Document()
    d.add_paragraph("Hello world.")
    tbl = d.add_table(rows=1, cols=2)
    tbl.rows[0].cells[0].text = "Auth Service"
    tbl.rows[0].cells[1].text = "The auth service shall authenticate users."
    d.save(str(path))


class HarvestAutoActorsTests(unittest.TestCase):
    def test_harvest_writes_sidecar_and_returns_path(self):
        with tempfile.TemporaryDirectory() as td:
            spec = Path(td) / "spec.docx"
            _write_minimal_docx(spec)
            out = Path(td) / "out.xlsx"

            args = SimpleNamespace(
                actors=None,
                nlp=False,
                config=None,
                keywords=None,
            )
            messages: list[str] = []
            auto_path = cli._harvest_auto_actors(
                args,
                [spec],
                out,
                progress=messages.append,
            )
            # Sidecar written next to the (would-be) requirements output.
            self.assertEqual(auto_path, out.with_name("out_auto_actors.xlsx"))
            self.assertTrue(auto_path.exists(), "sidecar .xlsx should be on disk")
            # At least one progress line mentioned harvesting.
            self.assertTrue(
                any("Auto-actors" in m for m in messages),
                "should announce the harvesting step via progress",
            )


if __name__ == "__main__":
    unittest.main()
