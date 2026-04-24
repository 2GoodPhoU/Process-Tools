"""Tests for the smart boilerplate-section auto-skip feature.

Defense / aerospace / telecoms specs all repeat the same handful of
boilerplate sections — Glossary, Acronyms, References, Revision
History, Document Control, etc. — that don't carry binding
requirements but do trigger the keyword classifier (e.g. a glossary
entry "shall mean" matches HARD).  Pre-skipping these by default
saves users from having to discover and configure each one per
project.

Implementation surface:

* ``config.DEFAULT_BOILERPLATE_TITLES`` — the built-in pattern list.
* ``config.SkipSections.auto_boilerplate`` — boolean toggle (defaults
  True so the skip is opt-out).
* ``config.SkipSections.matches_title()`` — extended to OR the default
  list when ``auto_boilerplate`` is True.
* ``parser._ParseContext.skip_heading_level`` — tracks heading-scope
  skip activated by a top-level boilerplate Heading.
* ``parser._emit_candidate`` — drops candidates while skip is active.
* ``parser.parse_docx_events`` — manages the heading-scope flag via
  a clear-then-set rule so adjacent boilerplate headings nest
  correctly.

Tests pin the title-matching contract first (cheap, pure functions)
then the heading-scope plumbing end-to-end via small synthetic docx
fixtures.
"""

from __future__ import annotations

import tempfile
import unittest
from pathlib import Path

from docx import Document

from requirements_extractor.actors import ActorResolver
from requirements_extractor.config import (
    DEFAULT_BOILERPLATE_TITLES,
    Config,
    SkipSections,
)
from requirements_extractor.models import RequirementEvent
from requirements_extractor.parser import parse_docx_events


# ---------------------------------------------------------------------------
# Pure-function: title matching
# ---------------------------------------------------------------------------


class TestDefaultBoilerplateTitles(unittest.TestCase):
    def test_includes_canonical_names(self):
        # The whole point of the default list — these have to be there
        # or the auto-skip stops being useful.
        for required in (
            "glossary", "definitions", "acronyms", "references",
            "revision history", "document control", "table of contents",
        ):
            self.assertIn(required, DEFAULT_BOILERPLATE_TITLES,
                          f"{required!r} missing from default list")

    def test_all_lowercase_for_case_insensitive_matching(self):
        # ``matches_title`` casefolds inputs before comparing — keep the
        # default list in lowercase so the comparison is symmetric and
        # the visual review of the list isn't deceived by case.
        for entry in DEFAULT_BOILERPLATE_TITLES:
            self.assertEqual(entry, entry.lower(),
                             f"{entry!r} not lowercase")


class TestMatchesTitleBoilerplateOn(unittest.TestCase):
    """auto_boilerplate=True (default) — the skip catches common names."""

    def setUp(self):
        self.skip = SkipSections()  # defaults

    def test_exact_match(self):
        self.assertTrue(self.skip.matches_title("Glossary"))
        self.assertTrue(self.skip.matches_title("References"))

    def test_case_insensitive(self):
        self.assertTrue(self.skip.matches_title("GLOSSARY"))
        self.assertTrue(self.skip.matches_title("revision HISTORY"))

    def test_numeric_prefix(self):
        # "3. Revision History", "Annex A — References", "12 Glossary" —
        # the substring rule from the original ``titles`` matcher
        # carries through to boilerplate too.
        self.assertTrue(self.skip.matches_title("3. Revision History"))
        self.assertTrue(self.skip.matches_title("Annex A — References"))
        self.assertTrue(self.skip.matches_title("12 Glossary"))

    def test_non_boilerplate_passes(self):
        # The whole reason the tool exists — substantive sections
        # must not get caught by the auto-skip.
        self.assertFalse(self.skip.matches_title("System Requirements"))
        self.assertFalse(self.skip.matches_title("3.2 Authentication"))
        self.assertFalse(self.skip.matches_title("Functional Behaviour"))

    def test_empty_title(self):
        self.assertFalse(self.skip.matches_title(""))
        self.assertFalse(self.skip.matches_title("   "))


class TestMatchesTitleBoilerplateOff(unittest.TestCase):
    """auto_boilerplate=False — only user ``titles`` apply."""

    def test_default_names_no_longer_match(self):
        skip = SkipSections(auto_boilerplate=False)
        self.assertFalse(skip.matches_title("Glossary"))
        self.assertFalse(skip.matches_title("References"))

    def test_user_titles_still_work(self):
        skip = SkipSections(titles=["Internal Notes"], auto_boilerplate=False)
        self.assertTrue(skip.matches_title("Internal Notes"))
        self.assertFalse(skip.matches_title("Glossary"))

    def test_user_titles_combine_with_boilerplate_when_on(self):
        skip = SkipSections(titles=["Internal Notes"], auto_boilerplate=True)
        self.assertTrue(skip.matches_title("Internal Notes"))
        self.assertTrue(skip.matches_title("Glossary"))


# ---------------------------------------------------------------------------
# YAML schema acceptance — the validator must allow the new key
# ---------------------------------------------------------------------------


class TestConfigValidatorAcceptsAutoBoilerplate(unittest.TestCase):
    def test_validator_accepts_auto_boilerplate_true(self):
        # Round-trip via the public load_config_raw machinery — uses
        # the same validator the CLI / GUI go through.
        from requirements_extractor.config import load_config_raw
        with tempfile.TemporaryDirectory() as d:
            cfg_path = Path(d) / "test.yaml"
            cfg_path.write_text(
                "skip_sections:\n  auto_boilerplate: false\n"
            )
            raw = load_config_raw(cfg_path)
            self.assertEqual(raw["skip_sections"]["auto_boilerplate"], False)


# ---------------------------------------------------------------------------
# End-to-end: heading-scope skip via a synthetic docx
# ---------------------------------------------------------------------------


def _build_doc_with_glossary(path: Path) -> None:
    """Spec layout: Heading 'System Requirements' (real) → table with
    one shall row, then Heading 'Glossary' → table with one shall row
    (boilerplate that should be skipped).
    """
    doc = Document()
    doc.add_heading("System Requirements", level=1)
    t1 = doc.add_table(rows=1, cols=2); t1.style = "Table Grid"
    t1.rows[0].cells[0].text = "Auth Service"
    t1.rows[0].cells[1].text = "The system shall verify all credentials."
    doc.add_heading("Glossary", level=1)
    t2 = doc.add_table(rows=1, cols=2); t2.style = "Table Grid"
    t2.rows[0].cells[0].text = "credential"
    # This text would otherwise classify as Hard via 'shall'.
    t2.rows[0].cells[1].text = "A credential shall mean any token issued by Auth."
    doc.save(str(path))


def _build_doc_with_glossary_then_requirements(path: Path) -> None:
    """Order: Glossary first, then real requirements section.  Skip
    must clear when the second top-level heading fires so the real
    requirement is captured."""
    doc = Document()
    doc.add_heading("Glossary", level=1)
    t1 = doc.add_table(rows=1, cols=2); t1.style = "Table Grid"
    t1.rows[0].cells[0].text = "token"
    t1.rows[0].cells[1].text = "A token shall mean an opaque string."
    doc.add_heading("System Requirements", level=1)
    t2 = doc.add_table(rows=1, cols=2); t2.style = "Table Grid"
    t2.rows[0].cells[0].text = "Auth Service"
    t2.rows[0].cells[1].text = "The system shall verify all credentials."
    doc.save(str(path))


def _extract_requirements(src: Path, *, auto_boilerplate: bool):
    cfg = Config.defaults()
    cfg.skip_sections.auto_boilerplate = auto_boilerplate
    resolver = ActorResolver()
    events = parse_docx_events(src, resolver.resolve, config=cfg)
    return [e.requirement for e in events if isinstance(e, RequirementEvent)]


class TestHeadingScopeSkipEndToEnd(unittest.TestCase):
    def test_glossary_after_real_section_drops_glossary_only(self):
        with tempfile.TemporaryDirectory() as d:
            src = Path(d) / "spec.docx"
            _build_doc_with_glossary(src)
            reqs = _extract_requirements(src, auto_boilerplate=True)
            # Only the real requirement survives; the glossary "shall"
            # row is dropped because it's under the Glossary heading.
            self.assertEqual(len(reqs), 1)
            self.assertIn("verify all credentials", reqs[0].text)

    def test_glossary_kept_when_auto_boilerplate_off(self):
        # Sanity: turning the switch off restores the pre-feature
        # behaviour where the glossary entries would slip into the
        # output.  This matters for users whose corpus uses one of
        # the default names for a non-boilerplate section.
        with tempfile.TemporaryDirectory() as d:
            src = Path(d) / "spec.docx"
            _build_doc_with_glossary(src)
            reqs = _extract_requirements(src, auto_boilerplate=False)
            texts = [r.text for r in reqs]
            self.assertEqual(len(reqs), 2)
            self.assertTrue(any("verify all credentials" in t for t in texts))
            self.assertTrue(any("opaque" in t or "credential shall mean" in t
                                for t in texts))

    def test_skip_clears_at_next_top_level_heading(self):
        # Glossary first → real Heading next.  The skip must clear at
        # the second H1 so the real requirement is captured.
        with tempfile.TemporaryDirectory() as d:
            src = Path(d) / "spec.docx"
            _build_doc_with_glossary_then_requirements(src)
            reqs = _extract_requirements(src, auto_boilerplate=True)
            self.assertEqual(len(reqs), 1)
            self.assertIn("verify all credentials", reqs[0].text)

    def test_section_row_title_path_also_skips(self):
        # The existing section-row title-match path (used inside 2-col
        # tables that don't have top-level headings) must continue to
        # honour auto-boilerplate.  Build a doc whose entire content
        # is one 2-col table where the column-1 text is "References".
        doc = Document()
        table = doc.add_table(rows=1, cols=2); table.style = "Table Grid"
        table.rows[0].cells[0].text = "References"
        table.rows[0].cells[1].text = "Section 5 shall apply to all systems."
        with tempfile.TemporaryDirectory() as d:
            src = Path(d) / "spec.docx"
            doc.save(str(src))
            reqs = _extract_requirements(src, auto_boilerplate=True)
            self.assertEqual(reqs, [])


if __name__ == "__main__":
    unittest.main()
